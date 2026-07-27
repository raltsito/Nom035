"""Anexo fotográfico (DOCX): muestra impresa de la evidencia de aplicación.

El documento NO sustituye a la galería: lleva deliberadamente solo una
selección de fotografías y explica en su primera página cómo consultar la
totalidad dentro de la plataforma. Reutiliza `plantilla_informe.docx` para
salir con el mismo membrete y estilos que el Informe Diagnóstico.
"""
import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from . import docx_geometry as geom

_PLANTILLA = Path(__file__).resolve().parent / 'plantilla_informe.docx'

FOTOS_POR_ANEXO = 20
COLUMNAS = 3
ANCHO_FOTO_CM = 4.6


def seleccionar_muestra(fotos, cantidad=FOTOS_POR_ANEXO):
    """Muestreo sistemático sobre la lista completa: toma cada k-ésima foto en
    vez de las primeras N, para que la muestra recorra toda la planta (áreas y
    números de empleado) y no solo el inicio del listado. Determinista: el
    mismo ciclo produce siempre el mismo anexo."""
    total = len(fotos)
    if total <= cantidad:
        return list(fotos)
    paso = total / cantidad
    return [fotos[min(int(i * paso), total - 1)] for i in range(cantidad)]


def build_anexo_fotografico_docx(planta, anio, fotos, total_fotos, muestra_informe,
                                 cobertura_pct, fecha_texto) -> io.BytesIO:
    """`fotos`: iterable de dicts {imagen: bytes, num_empleado, area, fecha}."""
    doc = Document(str(_PLANTILLA))

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run('ANEXO FOTOGRÁFICO')
    run.bold = True
    run.font.size = Pt(28)

    p_nom = doc.add_paragraph()
    p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_nom.add_run('NOM-035-STPS-2018')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xC8, 0x10, 0x2E)

    p_planta = doc.add_paragraph()
    p_planta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_planta.paragraph_format.space_before = Pt(10)
    run = p_planta.add_run(f'PLANTA: {planta}')
    run.bold = True
    run.font.size = Pt(13)

    p_ciclo = doc.add_paragraph()
    p_ciclo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ciclo.add_run(f'Ciclo de evaluación {anio} — generado el {fecha_texto}').font.size = Pt(11)

    doc.add_heading('Objeto de este anexo', level=1)
    doc.add_paragraph(
        'Este documento integra la evidencia fotográfica de la aplicación de la Guía de '
        'Referencia III de la NOM-035-STPS-2018 en el centro de trabajo. Cada fotografía fue '
        'capturada por el propio trabajador al inicio de su cuestionario, como constancia de '
        'que la aplicación se realizó de manera individual y presencial.'
    )

    doc.add_heading('Alcance: esta es una muestra', level=1)
    p_aviso = doc.add_paragraph()
    run = p_aviso.add_run(
        f'Este anexo contiene únicamente {len(fotos)} fotografías de las {total_fotos} '
        f'registradas para esta planta en el ciclo {anio}. Se trata de una muestra '
        'representativa —tomada de forma distribuida a lo largo de todo el padrón, no de los '
        'primeros registros— cuyo propósito es documentar la aplicación en un formato '
        'imprimible de extensión razonable.'
    )
    run.bold = True
    doc.add_paragraph(
        f'La cobertura fotográfica del ciclo es de {cobertura_pct}%: {total_fotos} de '
        f'{muestra_informe} trabajadores que integran la muestra válida del Informe '
        'Diagnóstico cuentan con fotografía. El resto de las imágenes no se reproduce aquí, '
        'pero está disponible en su totalidad dentro de la plataforma.'
    )

    doc.add_heading('Cómo consultar todas las fotografías', level=1)
    doc.add_paragraph(
        'La galería completa, con las fotografías de todos los trabajadores del ciclo, se '
        'consulta en línea siguiendo estos pasos:'
    )
    num_id = geom.nueva_lista(doc, 'decimal')
    pasos = [
        'Ingrese a la plataforma con su cuenta de administrador del centro de trabajo.',
        'Abra el módulo "Resultados" desde el menú lateral.',
        f'Seleccione el ciclo {anio} en el selector de la parte superior.',
        'Pulse el botón "Ver informe de fotografías", junto a "Descargar informe".',
        'Recorra la galería y pulse cualquier fotografía para verla en tamaño completo.',
    ]
    for paso in pasos:
        p = geom.parrafo_lista(doc, num_id)
        p.add_run(paso)

    p_conf = doc.add_paragraph()
    p_conf.add_run(
        'Confidencialidad: las fotografías y los datos que las acompañan son información '
        'personal de los trabajadores. Su consulta, reproducción y resguardo deben limitarse '
        'al personal autorizado del centro de trabajo, conforme al principio de '
        'confidencialidad de la NOM-035-STPS-2018.'
    ).italic = True

    doc.add_page_break()

    doc.add_heading('Muestra fotográfica', level=1)
    doc.add_paragraph(
        f'Se presentan {len(fotos)} fotografías identificadas con el número de empleado y el '
        'área de adscripción.'
    ).runs[0].italic = True

    filas = (len(fotos) + COLUMNAS - 1) // COLUMNAS
    if filas:
        table = doc.add_table(rows=filas, cols=COLUMNAS)
        table.style = 'Table Grid'
        for idx, foto in enumerate(fotos):
            celda = table.rows[idx // COLUMNAS].cells[idx % COLUMNAS]
            celda.text = ''
            p_img = celda.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(4)
            p_img.paragraph_format.space_after = Pt(2)
            try:
                p_img.add_run().add_picture(io.BytesIO(foto['imagen']), width=Cm(ANCHO_FOTO_CM))
            except Exception:
                p_img.add_run('(imagen no disponible)').font.size = Pt(8)
            p_pie = celda.add_paragraph()
            p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_pie.paragraph_format.space_after = Pt(6)
            r_num = p_pie.add_run(str(foto['num_empleado']))
            r_num.bold = True
            r_num.font.size = Pt(8)
            r_area = p_pie.add_run('\n' + str(foto['area']))
            r_area.font.size = Pt(7)
        # Celdas sobrantes de la última fila: vacías, sin borde de texto.
        for sobrante in range(len(fotos), filas * COLUMNAS):
            table.rows[sobrante // COLUMNAS].cells[sobrante % COLUMNAS].text = ''
        geom.fijar_geometria_tabla(table, anchos_twips=geom.anchos_equitativos(COLUMNAS),
                                   repetir_encabezado=False, dividir_filas=False)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
