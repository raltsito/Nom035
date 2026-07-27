"""Generador del borrador DOCX del reporte psicológico NOM-035.

Consume el mismo ctx que produce _build_psico_context() en views.py — nunca
importa interpretaciones.py ni contenido_normativo.py directamente. Esto
garantiza que el borrador y el HTML/PDF final queden idénticos en redacción
y estructura (las 13 secciones del estándar LEAR), sin duplicar lógica de
negocio en dos lugares.
"""
import io
from pathlib import Path

from django.contrib.staticfiles import finders
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

from . import contenido_informe as inf
from . import docx_geometry as geom

# Plantilla del Informe Diagnóstico (formato aprobado por dirección jul-2026):
# conserva estilos de Word (Grid Table 1 Light Accent 2, toc 1/3, etc.),
# encabezado con membrete, pie con número de página y updateFields para que
# Word reconstruya el índice al abrir. El cuerpo se genera aquí.
_PLANTILLA_INFORME = Path(__file__).resolve().parent / 'plantilla_informe.docx'

NIVEL_LABEL_DOCX = {
    'nulo': 'Nulo', 'bajo': 'Bajo', 'medio': 'Medio',
    'alto': 'Alto', 'muy_alto': 'Muy alto',
}


def _set_base_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)


def _heading(doc, text, level=2):
    doc.add_heading(text, level=level)


def _cell_bg(cell, hex6):
    """Colorea el fondo de una celda de tabla (mismo mecanismo que usa
    generar_informe.py para las tablas con encabezado de color)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    tcPr.append(shd)


def _cell_run(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _simple_table(doc, headers, rows, familia=None, anchos=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_vals in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row_vals):
            cells[i].text = str(v)
    if familia or anchos:
        geom.fijar_geometria_tabla(table, anchos_twips=anchos, familia=familia)
    doc.add_paragraph()


def _add_field(paragraph, instr_text, resultado=None):
    """Inserta un campo de Word (ej. ' PAGE ') en un paragraph, con la
    secuencia begin/instrText/separate/[resultado]/end — mismo mecanismo que
    usa `_add_indice` para el campo TOC.

    `resultado` (texto entre separate y end) es OBLIGATORIO para que
    LibreOffice importe un campo TOC como DocumentIndex actualizable: con el
    resultado vacío el importador de LO ignora el campo y el índice jamás se
    puede poblar (verificado con LibreOffice 24/25)."""
    partes = [('begin', None), (None, instr_text), ('separate', None)]
    if resultado is not None:
        partes.append(('texto', resultado))
    partes.append(('end', None))
    for ftype, text in partes:
        run = paragraph.add_run()
        if ftype == 'texto':
            run.text = text
        elif ftype:
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), ftype)
            run._r.append(fc)
        else:
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = text
            run._r.append(instr)


def _add_membrete(doc):
    """Encabezado y pie de página con el membrete de marca (banda
    roja/negra + logo LEAR), igual en todas las páginas — reproduce el
    documento estándar de formato (INFORME DIAGNOSTICO NOM035 LEAR
    TLÁHUAC.pdf). El número de página es un campo dinámico de Word, no
    parte de la imagen."""
    section = doc.sections[0]
    ancho_util = section.page_width - section.left_margin - section.right_margin

    # `header_distance`/`footer_distance` (espacio entre el borde de la
    # página y el encabezado/pie) venían por defecto en 1.27cm, más chico
    # que la altura real de las imágenes del membrete (~1.3-1.4cm) — Word
    # las recortaba contra el margen del cuerpo. Se reduce a 0.5cm para
    # dejarles espacio de sobra sin tocar los márgenes del cuerpo del texto.
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    header_path = finders.find('documents/img/membrete_header.png')
    if header_path:
        p = section.header.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(header_path, width=ancho_util)

    footer_path = finders.find('documents/img/membrete_footer_izq.png')
    if footer_path:
        p = section.footer.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(ancho_util, WD_TAB_ALIGNMENT.RIGHT)
        p.add_run().add_picture(footer_path, width=Cm(13))
        p.add_run('\t')
        _add_field(p, ' PAGE ')


def _add_sello(doc):
    """Sello NOM-035 de la portada (documento estándar de formato)."""
    sello_path = finders.find('documents/img/sello_nom035.png')
    if sello_path:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.add_run().add_picture(sello_path, width=Cm(4))


def _add_portada(doc, ctx):
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run('BORRADOR — Sujeto a revisión y aprobación profesional')
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x20, 0x20)
    run.font.size = Pt(12)

    doc.add_heading('Informe diagnóstico — Factores de Riesgo Psicosocial en el Trabajo', level=0)
    p = doc.add_paragraph()
    p.add_run('NOM-035-STPS-2018').bold = True
    p2 = doc.add_paragraph()
    p2.add_run('Centro de trabajo: ').bold = True
    p2.add_run(str(ctx['tenant'].nombre))
    p3 = doc.add_paragraph()
    p3.add_run('Ciclo de evaluación: ').bold = True
    p3.add_run(str(ctx['ciclo']))
    p4 = doc.add_paragraph()
    p4.add_run('Fecha de generación del borrador: ').bold = True
    p4.add_run(ctx['fecha_generado'])
    _add_sello(doc)
    doc.add_page_break()


_DOMICILIO_ETIQUETAS = ('Calle y No', 'Colonia', 'Municipio', 'Estado')


def _add_datos_centro_trabajo(doc, ctx):
    datos = ctx['datos_centro_trabajo']
    _heading(doc, '2. Datos del centro de trabajo', level=1)
    _heading(doc, '2.1. Centro de trabajo', level=3)
    p = doc.add_paragraph()
    p.add_run('Nombre, denominación o razón social: ').bold = True
    p.add_run(datos['nombre'])
    _heading(doc, '2.2. Domicilio', level=3)
    partes = (datos['direccion'] or '').split(' | ')
    if len(partes) == len(_DOMICILIO_ETIQUETAS):
        for etiqueta, valor in zip(_DOMICILIO_ETIQUETAS, partes):
            p = doc.add_paragraph()
            p.add_run(f'{etiqueta}: ').bold = True
            p.add_run(valor)
    else:
        doc.add_paragraph(datos['direccion'] or '—')
    _heading(doc, '2.3. Actividad principal', level=3)
    doc.add_paragraph(datos['giro'] or '—')


def _add_objetivo(doc, ctx):
    _heading(doc, '3. Objetivo', level=1)
    _heading(doc, '3.1. Objetivo general', level=3)
    doc.add_paragraph(ctx['objetivo_general'])
    _heading(doc, '3.2. Objetivos específicos', level=3)
    for obj in ctx['objetivos_especificos']:
        doc.add_paragraph(obj, style='List Number')


def _add_definiciones(doc, ctx):
    _heading(doc, '4. Definiciones', level=1)
    for termino, definicion in ctx['definiciones']:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{termino}: ').bold = True
        p.add_run(definicion)


def _add_justificacion_muestra(doc, ctx):
    jm = ctx['justificacion_muestra']
    _heading(doc, '5. Justificación de la muestra', level=1)

    if jm.get('parrafos'):
        for parrafo in jm['parrafos']:
            doc.add_paragraph(parrafo)
    else:
        doc.add_paragraph(jm['intro'])
        doc.add_paragraph(
            f"La población del centro de trabajo es de {jm['poblacion']} trabajadores; se evaluó "
            f"una muestra de {jm['muestra']} participantes, de los cuales {jm['hombres']} fueron "
            f"hombres y {jm['mujeres']} mujeres."
        )

    eq1 = jm.get('ecuacion1')
    if not eq1:
        return

    def _pct(parte, total):
        return f'{parte / total * 100:.1f}%' if total else '0.0%'

    table = doc.add_table(rows=3, cols=6)
    table.style = 'Table Grid'

    encabezados = ['Colaboradores', 'Hombres', 'Mujeres', 'Muestra', 'Hombres', 'Mujeres']
    for j, h in enumerate(encabezados):
        cell = table.rows[0].cells[j]
        _cell_bg(cell, 'C8102E')
        _cell_run(cell, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))

    subencabezados = ['Total', '(N / %)', '(N / %)', f"N mín. {eq1['n_min']}", '(N / %)', '(N / %)']
    for j, s in enumerate(subencabezados):
        cell = table.rows[1].cells[j]
        _cell_bg(cell, '1A1A2E')
        _cell_run(cell, s, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))

    datos = [
        str(eq1['total']),
        f"{eq1['hombres']} / {_pct(eq1['hombres'], eq1['total'])}",
        f"{eq1['mujeres']} / {_pct(eq1['mujeres'], eq1['total'])}",
        str(eq1['muestra']),
        f"{eq1['muestra_hombres']} / {_pct(eq1['muestra_hombres'], eq1['muestra'])}",
        f"{eq1['muestra_mujeres']} / {_pct(eq1['muestra_mujeres'], eq1['muestra'])}",
    ]
    for j, d in enumerate(datos):
        _cell_run(table.rows[2].cells[j], d, size=10)

    doc.add_paragraph()

    if ctx.get('flujo_muestra'):
        doc.add_paragraph(
            'Flujo de la muestra (Guía III) — de la población total a los '
            'cuestionarios analizados:'
        ).runs[0].italic = True
        _simple_table(
            doc,
            ['Etapa', 'N'],
            [[f['etapa'], f['n']] for f in ctx['flujo_muestra']],
        )
        doc.add_paragraph()


def _add_metodologia(doc, ctx):
    met = ctx['metodologia']
    _heading(doc, '6. Metodología', level=1)
    _heading(doc, '6.1. Objetivo de la evaluación', level=3)
    doc.add_paragraph(met['objetivo_evaluacion'])

    _heading(doc, '6.2. Instrumentos utilizados', level=3)
    _simple_table(doc, ['Instrumento', 'Descripción'], met['instrumentos'])

    _heading(doc, '6.3. Procedimiento de aplicación', level=3)
    doc.add_paragraph('Antes de la aplicación:', style='Intense Quote')
    for paso in met['procedimiento_antes']:
        doc.add_paragraph(paso, style='List Number')
    doc.add_paragraph('Durante la aplicación del cuestionario:', style='Intense Quote')
    for paso in met['procedimiento_durante']:
        doc.add_paragraph(paso, style='List Number')
    doc.add_paragraph('Después de la aplicación del cuestionario:', style='Intense Quote')
    for paso in met['procedimiento_despues']:
        doc.add_paragraph(paso, style='List Number')

    _heading(doc, '6.4. Procesamiento a digital', level=3)
    doc.add_paragraph(met['procesamiento_digital'])

    _heading(doc, '6.5. Evaluación y análisis de resultados', level=3)
    doc.add_paragraph(met['evaluacion_analisis'])

    if ctx['tasas_respuesta']:
        _heading(doc, 'Tasa de respuesta por instrumento', level=3)
        _simple_table(
            doc,
            ['Instrumento', 'Aplicados', 'Respondidos', 'Tasa'],
            [[t['guia'], t['total'], t['completadas'], f"{t['pct']}%"] for t in ctx['tasas_respuesta']],
        )


def _add_grafica(doc, png_bytesio, width_cm=9):
    """Embebe una gráfica (BytesIO de PNG) centrada, si existe. No-op si
    `png_bytesio` es None (borrador sin `informe_extendido` nunca las pasa)."""
    if not png_bytesio:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(png_bytesio, width=Cm(width_cm))


def _add_informe_demografico(doc, ctx):
    _heading(doc, '7. Informe demográfico', level=1)
    doc.add_paragraph(
        f"El perfil demográfico corresponde a la población analítica del informe: "
        f"{ctx['muestra']['total']} trabajadores con cuestionario de la Guía III válido. "
        f"Cada variable reporta su N válido; los porcentajes se calculan sobre ese denominador."
    )
    rd = ctx.get('report_data')
    if rd and rd['faltantes_criticos']['variables']:
        doc.add_paragraph(
            'Alerta: las variables siguientes superan el umbral de datos faltantes '
            f"({rd['faltantes_criticos']['umbral_pct']}%) y se presentan sin gráfica: "
            + '; '.join(f"{v['variable']} ({v['pct_faltante']}% sin dato)"
                        for v in rd['faltantes_criticos']['variables'])
        ).runs[0].italic = True
    graficas_muestra = ctx.get('graficas', {}).get('muestra', {})

    _heading(doc, '7.1. Informe de datos generales', level=3)
    cols_generales = {'Sexo', 'Grupo de edad', 'Estado civil', 'Nivel de estudios'}
    for tabla in ctx['muestra_tablas']:
        if tabla['col'] in cols_generales:
            _heading(doc, tabla['titulo'], level=4)
            _simple_table(
                doc, [tabla['col'], 'Trabajadores', '%'],
                [[d['label'], d['count'], f"{d['pct']}%"] for d in tabla['datos']],
            )
            _add_grafica(doc, graficas_muestra.get(tabla['col']))

    _heading(doc, '7.1.1. Informe de datos laborales', level=3)
    for tabla in ctx['muestra_tablas']:
        if tabla['col'] not in cols_generales:
            _heading(doc, tabla['titulo'], level=4)
            _simple_table(
                doc, [tabla['col'], 'Trabajadores', '%'],
                [[d['label'], d['count'], f"{d['pct']}%"] for d in tabla['datos']],
            )
            _add_grafica(doc, graficas_muestra.get(tabla['col']))


def _add_ats_tabla(doc, filas):
    """Tabla de desglose por pregunta: # | Pregunta | H-Sí | H-No | M-Sí |
    M-No | Total-Sí | Total-No — mismo formato que generar_informe.py."""
    encabezados = ['#', 'Pregunta', 'H\nSí', 'H\nNo', 'M\nSí', 'M\nNo', 'Total\nSí', 'Total\nNo']
    table = doc.add_table(rows=1, cols=len(encabezados))
    table.style = 'Table Grid'
    for j, h in enumerate(encabezados):
        cell = table.rows[0].cells[j]
        _cell_bg(cell, 'C8102E')
        _cell_run(cell, h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    for f in filas:
        cells = table.add_row().cells
        valores = [f['orden'], f['texto'], f['h_si'], f['h_no'], f['m_si'], f['m_no'], f['t_si'], f['t_no']]
        for j, v in enumerate(valores):
            _cell_run(cells[j], v, size=9, align=(WD_ALIGN_PARAGRAPH.LEFT if j == 1 else WD_ALIGN_PARAGRAPH.CENTER))
    geom.fijar_geometria_tabla(table, familia='ats_desglose_8')
    doc.add_paragraph()


def _add_ats(doc, ctx):
    guia_i = ctx['guia_i']
    _heading(doc, '8. Informe de acontecimientos traumáticos severos (ATS)', level=1)
    doc.add_paragraph(ctx['guia_i_definicion'])
    doc.add_paragraph(ctx['guia_i_texto'])
    if guia_i['total'] == 0:
        doc.add_paragraph('No se aplicó la Guía de Referencia I en este ciclo.')
        return

    desglose = ctx.get('ats_desglose')
    if not desglose:
        doc.add_paragraph(
            f"De los {guia_i['total']} trabajadores que respondieron la Guía I, "
            f"{guia_i['requieren_atencion']} cumplen el criterio de caso positivo."
        )
        if guia_i['casos']:
            _simple_table(
                doc,
                ['Trabajador', 'Área', 'Acontecimiento (D1)', 'Síntomas (D2–D4)'],
                [[c['trabajador_nombre'], c['trabajador_area'], c['acontecimiento'], c['sintomas']]
                 for c in guia_i['casos']],
            )
        return

    rd = ctx.get('report_data')
    if rd:
        # Resumen oficial desde ReportDataNOM035 (fuente única).
        ats = rd['tablas']['ats']
        doc.add_paragraph(ats['nota']).runs[0].italic = True
        _simple_table(
            doc,
            ['Indicador', 'N', '%'],
            [[f['indicador'], f['n'], f"{f['pct']}%" if f['pct'] is not None else '—']
             for f in ats['filas']],
        )
        doc.add_paragraph(
            f"Denominador: {ats['denominador']} (N = {ats['n_valido']}). "
            'El desglose por reactivo se presenta en el Anexo técnico.'
        ).runs[0].italic = True
    else:
        r = desglose['resumen']
        doc.add_paragraph(
            f"De los {r['muestra']} trabajadores que respondieron la Guía I, {r['con_ats']} "
            f"({r['pct_ats']}%) reportaron al menos un acontecimiento traumático severo, y "
            f"{r['clinica']} ({r['pct_clinica']}%) cumplen el criterio de valoración clínica."
        )
        _simple_table(
            doc,
            ['Colaboradores evaluados', 'Reportaron ATS', '% ATS', 'Requieren valoración clínica', '% valoración clínica'],
            [[r['muestra'], r['con_ats'], f"{r['pct_ats']}%", r['clinica'], f"{r['pct_clinica']}%"]],
        )


def _fmt_pct(v):
    return f'{v}%' if v is not None else '—'


def _tabla_niveles_rd(doc, tabla_rd, col1):
    """Tabla estándar de niveles (categorías/dominios) desde ReportDataNOM035:
    N por nivel con %, acumulados "Medio o superior" (referencia para el
    Programa de intervención) y "Alto o Muy alto" (prioridad elevada), y la
    prioridad del ranking. Filas en orden de prioridad."""
    filas = sorted(tabla_rd['filas'], key=lambda f: f['prioridad'])
    _simple_table(
        doc,
        [col1, 'N', 'Nulo %', 'Bajo %', 'Medio %', 'Alto %', 'Muy alto %',
         'Medio o superior %', 'Alto o Muy alto %', 'Prioridad'],
        [[f['nombre'], f['n']]
         + [_fmt_pct(f['pct_nivel'][k]) for k in ('nulo', 'bajo', 'medio', 'alto', 'muy_alto')]
         + [_fmt_pct(f['pct_medio_o_mas']), _fmt_pct(f['pct_alto_o_muy_alto']), f"{f['prioridad']}°"]
         for f in filas],
    )
    doc.add_paragraph(
        f"Denominador: {tabla_rd['denominador']}. {tabla_rd['nota']}"
    ).runs[0].italic = True


def _add_panel_ejecutivo(doc, ctx):
    """Panel ejecutivo: tarjetas generadas por el motor de composición
    (ReportDataNOM035.tarjetas). Solo aparecen tarjetas con datos válidos;
    no existe tarjeta de 'nivel global organizacional'."""
    rd = ctx.get('report_data')
    if not rd or not rd.get('tarjetas'):
        return
    _heading(doc, 'Panel ejecutivo', level=1)
    doc.add_paragraph(
        'Indicadores clave del ciclo, derivados de un objeto único de datos '
        '(ReportDataNOM035) validado antes de la emisión. Los niveles de riesgo '
        'se clasifican por cuestionario individual; no se presenta ningún nivel '
        'global derivado del promedio.'
    ).runs[0].italic = True

    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    tarjetas = rd['tarjetas']
    for i in range(0, len(tarjetas), 3):
        row = table.add_row().cells
        for j, t in enumerate(tarjetas[i:i + 3]):
            _cell_bg(row[j], 'F0F0F0')
            p = row[j].paragraphs[0]
            r1 = p.add_run(t['titulo'] + '\n')
            r1.bold = True
            r1.font.size = Pt(9)
            r2 = p.add_run(str(t['valor']) + '\n')
            r2.bold = True
            r2.font.size = Pt(16)
            r2.font.color.rgb = RGBColor(0xC8, 0x10, 0x2E)
            if t.get('detalle'):
                r3 = p.add_run(t['detalle'])
                r3.font.size = Pt(8)
            if t.get('nota'):
                r4 = p.add_run('\n' + t['nota'])
                r4.font.size = Pt(7)
                r4.italic = True
    geom.fijar_geometria_tabla(table, familia='panel_3',
                               repetir_encabezado=False, dividir_filas=False)

    # (La línea "Validación de consistencia previa a la emisión…" se retiró
    # del formato aprobado: el maestro de Zapotitlán no la incluye. La
    # validación sigue ejecutándose y bloqueando la emisión en views.py.)

    if rd['faltantes_criticos']['variables']:
        doc.add_paragraph(
            'Alerta de faltantes críticos (no se grafican estas variables): '
            + '; '.join(f"{v['variable']} ({v['pct_faltante']}% sin dato)"
                        for v in rd['faltantes_criticos']['variables'])
        ).runs[0].italic = True
    doc.add_page_break()


def _add_informe_diagnostico(doc, ctx):
    _heading(doc, '9. Informe diagnóstico sobre los factores de riesgo psicosocial y del entorno organizacional', level=1)

    _heading(doc, '9.1. Calificación final', level=3)
    doc.add_paragraph(
        f"Nivel de riesgo predominante en la organización: {ctx['nivel_global_label']} "
        f"(moda de los niveles individuales; N = {sum(d['count'] for d in ctx['distribucion'])}). "
        f"El {ctx.get('pct_medio_o_mas', 0)}% de las personas evaluadas se ubica en nivel "
        f"Medio o superior y el {ctx.get('pct_alto_muy_alto', 0)}% en Alto o Muy alto. "
        f"Puntaje promedio: {ctx['promedio_global']}; mediana: {ctx.get('mediana_global', '—')}."
    )
    if ctx.get('nota_promedio'):
        doc.add_paragraph(ctx['nota_promedio']).runs[0].italic = True
    doc.add_paragraph(ctx['nivel_global_texto'])
    _simple_table(
        doc,
        ['Nivel de riesgo', 'Trabajadores', '%'],
        [[d['label'], d['count'], f"{d['pct']}%"] for d in ctx['distribucion']],
    )
    _add_grafica(doc, ctx.get('graficas', {}).get('distribucion'), width_cm=11)

    if ctx.get('acciones'):
        doc.add_paragraph(
            'Tabla 7, Guía de Referencia III, NOM-035-STPS-2018 — Criterios para la toma '
            'de acciones (transcripción del texto oficial):'
        ).runs[0].italic = True
        _simple_table(
            doc,
            ['Nivel de riesgo', 'Acción requerida'],
            [[a['label'], a['accion']] for a in ctx['acciones']],
        )

    graficas_categorias = ctx.get('graficas', {}).get('categorias', {})
    graficas_dominios = ctx.get('graficas', {}).get('dominios', {})

    _heading(doc, '9.2. Calificaciones por categoría', level=3)
    if ctx.get('rangos_categoria'):
        _niveles = ('nulo', 'bajo', 'medio', 'alto', 'muy_alto')
        doc.add_paragraph('Rangos oficiales de corte (Tabla 6, Guía de Referencia III):').runs[0].italic = True
        _simple_table(
            doc,
            ['Categoría'] + [NIVEL_LABEL_DOCX[n] for n in _niveles],
            [[r['nombre']] + [r['rangos'][n] for n in _niveles] for r in ctx['rangos_categoria']],
        )
    if ctx.get('report_data'):
        _tabla_niveles_rd(doc, ctx['report_data']['tablas']['categorias'], 'Categoría')
    else:
        _simple_table(
            doc,
            ['Categoría', 'Nulo', 'Bajo', 'Medio', 'Alto', 'Muy alto',
             'Alto o Muy alto (prioridad elevada)', '%'],
            [[c['nombre'], c['dist']['nulo'], c['dist']['bajo'], c['dist']['medio'],
              c['dist']['alto'], c['dist']['muy_alto'], c['requieren_atencion'], f"{c['pct_intervencion']}%"]
             for c in ctx['categorias_agregadas']],
        )
    for c in ctx['categorias_agregadas']:
        _add_grafica(doc, graficas_categorias.get(c['nombre']), width_cm=8)

    _heading(doc, '9.3. Calificaciones por dominio', level=3)
    if ctx.get('rangos_dominio'):
        _niveles = ('nulo', 'bajo', 'medio', 'alto', 'muy_alto')
        doc.add_paragraph('Rangos oficiales de corte (Tabla 6, Guía de Referencia III):').runs[0].italic = True
        _simple_table(
            doc,
            ['Dominio'] + [NIVEL_LABEL_DOCX[n] for n in _niveles],
            [[r['nombre']] + [r['rangos'][n] for n in _niveles] for r in ctx['rangos_dominio']],
        )
    if ctx.get('report_data'):
        _tabla_niveles_rd(doc, ctx['report_data']['tablas']['dominios'], 'Dominio')
    else:
        _simple_table(
            doc,
            ['Dominio', 'Nulo', 'Bajo', 'Medio', 'Alto', 'Muy alto',
             'Alto o Muy alto (prioridad elevada)', '%'],
            [[d['nombre'], d['dist']['nulo'], d['dist']['bajo'], d['dist']['medio'],
              d['dist']['alto'], d['dist']['muy_alto'], d['requieren_atencion'], f"{d['pct_intervencion']}%"]
             for d in ctx['dominios_oficiales_agregados']],
        )
    for d in ctx['dominios_oficiales_agregados']:
        _add_grafica(doc, graficas_dominios.get(d['nombre']), width_cm=8)
    for dom in ctx['dominios_prioritarios']:
        _heading(doc, f"{dom['clave']} — {dom['nombre']}", level=4)
        if dom.get('mide'):
            p = doc.add_paragraph()
            p.add_run('Qué mide: ').italic = True
            p.add_run(dom['mide'])
        if dom.get('interpretacion'):
            p = doc.add_paragraph()
            p.add_run(dom['interpretacion']).italic = True

    _heading(doc, '9.4. Dimensiones (indicador descriptivo)', level=3)
    if ctx.get('nota_dimensiones'):
        doc.add_paragraph(ctx['nota_dimensiones']).runs[0].italic = True
    if ctx['dimensiones'] and 'pct_promedio' in ctx['dimensiones'][0]:
        # Las 25 dimensiones oficiales de la Tabla 6 — sin niveles de riesgo
        # (la NOM no define cortes por dimensión).
        _simple_table(
            doc,
            ['#', 'Dimensión', 'Dominio', 'N', '% promedio', '% mediana'],
            [[dim['orden'], dim['nombre'], dim['dominio_oficial'], dim['n'],
              f"{dim['pct_promedio']}%", f"{dim['pct_mediana']}%"]
             for dim in ctx['dimensiones']],
        )
    rd = ctx.get('report_data')
    if not rd:
        return

    # ---- 9.5. Áreas prioritarias (distribución de niveles individuales) ----
    areas = rd['tablas']['areas']
    _heading(doc, '9.5. Áreas prioritarias', level=3)
    doc.add_paragraph(f"{areas['nota']} Denominador: {areas['denominador']}.").runs[0].italic = True
    if areas['filas']:
        _simple_table(
            doc,
            ['Área', 'N válido', 'Nulo %', 'Bajo %', 'Medio %', 'Alto %', 'Muy alto %',
             'Medio o superior %', 'Alto o Muy alto %', 'Prioridad'],
            [[a['nombre'], a['n']]
             + [_fmt_pct(a['pct_nivel'][k]) for k in ('nulo', 'bajo', 'medio', 'alto', 'muy_alto')]
             + [_fmt_pct(a['pct_medio_o_mas']), _fmt_pct(a['pct_alto_o_muy_alto']), f"{a['prioridad']}°"]
             for a in areas['filas']],
        )
    else:
        doc.add_paragraph('No hay áreas con N suficiente para el análisis desagregado.')
    if areas['reservadas']:
        doc.add_paragraph(
            'Áreas reservadas por confidencialidad (N menor al umbral de '
            f"{areas['umbral_confidencialidad']}): "
            + ', '.join(f"{a['nombre']} (N={a['n']})" for a in areas['reservadas'])
        ).runs[0].italic = True

    # ---- 9.6. Violencia laboral (resumen del dominio oficial) ----
    viol = rd['tablas']['violencia']
    _heading(doc, '9.6. Violencia laboral (resumen)', level=3)
    doc.add_paragraph(viol['nota']).runs[0].italic = True
    _simple_table(
        doc,
        ['Nivel', 'N', '%'],
        [[f['label'], f['n'], _fmt_pct(f['pct'])] for f in viol['filas']],
    )
    doc.add_paragraph(
        f"Denominador: {viol['denominador']} (N = {viol['n_valido']}). "
        f"Personal en Alto o Muy alto: {_fmt_pct(viol['pct_alto_o_muy_alto'])}."
    ).runs[0].italic = True


def _add_conclusiones_tabla_rankeada(doc, encabezado_col1, filas, familia=None,
                                     diseno_informe=False):
    """`diseno_informe=True` viste la tabla con el diseño aprobado del Informe
    Diagnóstico (encabezado rojo, banding y bordes finos) conservando las
    negritas propias de esta tabla — nivel predominante y prioridad. El
    borrador psicológico la sigue emitiendo con encabezado negro."""
    encabezados = [encabezado_col1, 'Nivel predominante', '% Programa de intervención (M+A+MA)', '% Alto+Muy alto', 'Prioridad']
    table = doc.add_table(rows=1, cols=len(encabezados))
    table.style = 'Table Grid'
    for j, h in enumerate(encabezados):
        cell = table.rows[0].cells[j]
        _cell_bg(cell, _TABLA_ENCABEZADO if diseno_informe else '1A1A2E')
        _cell_run(cell, h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
    for idx, g in enumerate(filas, 1):
        cells = table.add_row().cells
        _cell_run(cells[0], g['nombre'], size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_run(cells[1], NIVEL_LABEL_DOCX.get(g['categoria_predominante'], g['categoria_predominante']), bold=True, size=10)
        _cell_run(cells[2], f"{g['pct_accion']}%", size=10)
        _cell_run(cells[3], f"{g['pct_intervencion']}%", size=10)
        _cell_run(cells[4], f'{idx}°', bold=True, size=10)
    if familia:
        geom.fijar_geometria_tabla(table, familia=familia)
    if diseno_informe:
        _fijar_bordes_tabla(table)
        for i, row in enumerate(table.rows[1:]):
            if i % 2 == 1:
                for cell in row.cells:
                    _cell_bg(cell, _TABLA_BANDA)
        # Semáforo en la columna "Nivel predominante" (§10.2 y §10.3).
        for row, g in zip(table.rows[1:], filas):
            _pintar_celda_nivel(row.cells[1], g['categoria_predominante'],
                                align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


def _add_conclusiones(doc, ctx):
    c = ctx['conclusiones']
    _heading(doc, '10. Conclusiones', level=1)

    if not ctx.get('acciones_generales'):
        # Borrador: mismo texto de siempre, sin cambios.
        doc.add_paragraph(
            f"Conforme a la información analizada la calificación final, de una muestra "
            f"representativa de {c['muestra']} trabajadores, de un total de {c['poblacion']} "
            f"trabajadores; el {c['pct_riesgo']}% ({c['riesgo_alto']} trabajadores) perciben que "
            f"las condiciones de trabajo se tienen que mejorar."
        )
        if c['categorias_destacadas']:
            doc.add_paragraph(
                'De acuerdo a la información analizada en la agrupación correspondiente a las '
                'categorías, el personal presenta la mayor ponderación en:'
            )
            for cat in c['categorias_destacadas']:
                doc.add_paragraph(f"{cat['nombre']} — {cat['pct_intervencion']}%", style='List Bullet')
        if c['dominios_destacados']:
            doc.add_paragraph(
                'Por otra parte, en la agrupación correspondiente a los dominios, el personal '
                'presenta la mayor ponderación en:'
            )
            for dom in c['dominios_destacados']:
                doc.add_paragraph(f"{dom['nombre']} — {dom['pct_intervencion']}%", style='List Bullet')
        return

    doc.add_paragraph(
        f"Con base en los resultados obtenidos de la aplicación de la Guía de Referencia III "
        f"de la NOM-035-STPS-2018, con una muestra de {c['muestra']} colaboradores evaluados, "
        f"se presentan las siguientes conclusiones. Conforme al apartado III.4 de la Guía III, "
        f"el nivel de riesgo se determina a partir de la calificación de cada cuestionario, y "
        f"las acciones a adoptar se establecen con base en la Tabla 7 de la propia guía, "
        f"mediante un Programa de intervención para los niveles Medio, Alto y Muy alto."
    )

    _heading(doc, '10.1. Calificación final', level=3)
    doc.add_paragraph(
        f"De los {c['muestra']} trabajadores evaluados, {c['accion_global']} ({c['pct_accion']}%) "
        f"obtuvieron una calificación final en los niveles Medio, Alto o Muy alto, por lo que "
        f"conforme a la Tabla 7 del apartado III.4 de la Guía III requieren la adopción de "
        f"acciones para el control de los factores de riesgo psicosocial a través de un "
        f"Programa de intervención. De ellos, {c['riesgo_alto']} trabajadores ({c['pct_riesgo']}%) "
        f"se encuentran en los niveles Alto o Muy alto. El nivel de riesgo predominante en la "
        f"muestra (dato descriptivo) es "
        f"{NIVEL_LABEL_DOCX.get(c['nivel_predominante'], c['nivel_predominante']).upper()}."
    )
    _simple_table(
        doc,
        ['Muestra', 'Requieren Programa de intervención (M+A+MA)', 'Prioridad interna (Alto+Muy alto)', 'Nivel predominante'],
        [[c['muestra'], f"{c['accion_global']} ({c['pct_accion']}%)",
          f"{c['riesgo_alto']} ({c['pct_riesgo']}%)",
          NIVEL_LABEL_DOCX.get(c['nivel_predominante'], c['nivel_predominante'])]],
    )

    _heading(doc, '10.2. Análisis por categoría', level=3)
    doc.add_paragraph(
        'A continuación se presentan las cinco categorías evaluadas, ordenadas de mayor a '
        'menor porcentaje de población en niveles Alto y Muy alto (indicador interno de '
        'priorización).'
    )
    _add_conclusiones_tabla_rankeada(doc, 'Categoría', c['categorias_rankeadas'])

    _heading(doc, '10.3. Análisis por dominio', level=3)
    doc.add_paragraph(
        'Los dominios se presentan ordenados de mayor a menor porcentaje de población en '
        'niveles Alto y Muy alto, permitiendo identificar los aspectos específicos del entorno '
        'laboral que concentran mayor riesgo psicosocial.'
    )
    _add_conclusiones_tabla_rankeada(doc, 'Dominio', c['dominios_rankeados'])

    cat_top = c['categorias_rankeadas'][0]['nombre'] if c['categorias_rankeadas'] else '—'
    dom_top = c['dominios_rankeados'][0]['nombre'] if c['dominios_rankeados'] else '—'
    doc.add_paragraph(
        f"En síntesis, el {c['pct_accion']}% de los trabajadores evaluados obtuvo una "
        f"calificación final en niveles Medio, Alto o Muy alto, por lo que se requiere la "
        f"adopción de acciones mediante un Programa de intervención, conforme a la Tabla 7 del "
        f"apartado III.4 de la Guía III. La categoría que concentra el mayor porcentaje de "
        f"trabajadores en niveles Alto y Muy alto es \"{cat_top}\", y el dominio con mayor "
        f"prioridad de atención es \"{dom_top}\". Se recomienda implementar las acciones "
        f"descritas en la sección 11 de este informe, priorizando las áreas identificadas."
    )


def _add_recomendaciones(doc, ctx):
    _heading(doc, '11. Recomendaciones', level=1)

    if ctx.get('acciones_generales'):
        _heading(doc, '11.1. Recomendaciones generales', level=3)
        _simple_table(
            doc, ['Acción', 'Descripción'],
            [[titulo, texto] for titulo, texto in ctx['acciones_generales']],
        )

    recs_extendidas = ctx.get('recomendaciones_extendidas')
    if recs_extendidas is not None:
        _heading(doc, '11.2. Recomendaciones específicas', level=3)
        if not recs_extendidas:
            doc.add_paragraph(
                'Ningún dominio alcanzó nivel de riesgo Alto o Muy alto a nivel agregado. '
                'Se recomienda mantener las condiciones organizacionales actuales y reaplicar '
                'la evaluación en el siguiente ciclo normativo.'
            )
            return
        for idx, rec in enumerate(recs_extendidas, 1):
            _heading(doc, f"{idx}. {rec['dominio']} ({rec['pct_intervencion']}%)", level=4)
            doc.add_paragraph(rec['intro'])
            for titulo, texto in rec['bullets']:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f'{titulo}: ').bold = True
                p.add_run(texto)
        return

    _heading(doc, '11.1. Recomendaciones específicas', level=3)
    if not ctx['recomendaciones']:
        doc.add_paragraph(
            'Ningún dominio alcanzó nivel de riesgo Alto o Muy alto a nivel agregado. '
            'Se recomienda mantener las condiciones organizacionales actuales y reaplicar '
            'la evaluación en el siguiente ciclo normativo.'
        )
        return
    for rec in ctx['recomendaciones']:
        p = doc.add_paragraph(style='List Bullet')
        if rec.get('urgencia'):
            p.add_run(f"[{rec['urgencia']}] ").bold = True
        p.add_run(rec['texto'])
        if rec.get('dominio'):
            p.add_run(f"  ({rec['dominio']})").italic = True


def _add_datos_responsable(doc, ctx):
    _heading(doc, '12. Datos del responsable de la evaluación', level=1)
    # En el borrador estos datos siempre quedan en blanco: la identidad del
    # revisor solo se inyecta al momento de aprobar (ver _inyectar_validacion_revisor).
    doc.add_paragraph('Nombre: ______________________________')
    doc.add_paragraph('Número de cédula profesional: ______________________________')
    doc.add_paragraph('Fecha de aprobación: ______________________________')


def _add_anexos(doc, ctx, con_geometria=False):
    # Anexo confidencial (resultados individuales por folio): solo se emite
    # si el contexto lo incluyó expresamente (incluir_anexo_confidencial).
    if not ctx.get('anexos'):
        return
    anexos = ctx['anexos']
    _heading(doc, '13. Anexos', level=1)

    def _anchos(n_cols, primera=0.12):
        return geom.anchos_equitativos(n_cols, primera=primera) if con_geometria else None

    _heading(doc, '13.1. ATS y Calificación final', level=3)
    if ctx.get('acciones_generales'):
        _simple_table(
            doc,
            ['Folio', 'ATS', 'Valoración clínica', 'Calificación final', 'Nivel'],
            [[f['clave'], f['ats'], f['valoracion_clinica'], f['calificacion_final'],
              NIVEL_LABEL_DOCX.get(f['nivel_final'], f['nivel_final'])]
             for f in anexos['ats_final']],
            anchos=_anchos(5),
        )
    else:
        _simple_table(
            doc,
            ['Folio', 'Guía I — ATS', 'Calificación final', 'Nivel'],
            [[f['clave'], f['ats_combinado'], f['calificacion_final'],
              NIVEL_LABEL_DOCX.get(f['nivel_final'], f['nivel_final'])]
             for f in anexos['ats_final']],
            anchos=_anchos(4),
        )

    _heading(doc, '13.2. Categorías', level=3)
    _simple_table(
        doc,
        ['Folio'] + anexos['categorias_columnas'],
        [[f['clave']] + [NIVEL_LABEL_DOCX.get(n, '—') if n else '—' for n in f['niveles']]
         for f in anexos['categorias_filas']],
        anchos=_anchos(1 + len(anexos['categorias_columnas'])),
    )

    _heading(doc, '13.3. Dominios', level=3)
    _simple_table(
        doc,
        ['Folio'] + anexos['dominios_columnas'],
        [[f['clave']] + [NIVEL_LABEL_DOCX.get(n, '—') if n else '—' for n in f['niveles']]
         for f in anexos['dominios_filas']],
        anchos=_anchos(1 + len(anexos['dominios_columnas']), primera=0.08),
    )

    _heading(doc, '13.4. Tabla de agrupación de dominios', level=3)
    _simple_table(
        doc,
        ['Categoría', 'Dominio', 'Dimensiones (D1-D14)'],
        [[f['categoria'], f['dominio'], ', '.join(f['claves'])] for f in anexos['agrupacion']],
        anchos=_anchos(3, primera=0.30),
    )


def build_reporte_psicologico_docx(ctx: dict) -> io.BytesIO:
    """Construye el borrador DOCX a partir del ctx de _build_psico_context(),
    siguiendo la misma estructura de 13 secciones que el PDF final aprobado."""
    doc = Document()
    _set_base_styles(doc)
    _add_membrete(doc)

    _add_portada(doc, ctx)
    _add_datos_centro_trabajo(doc, ctx)
    _add_objetivo(doc, ctx)
    _add_definiciones(doc, ctx)
    _add_justificacion_muestra(doc, ctx)
    _add_metodologia(doc, ctx)
    _add_informe_demografico(doc, ctx)
    _add_ats(doc, ctx)
    _add_informe_diagnostico(doc, ctx)
    _add_conclusiones(doc, ctx)
    _add_recomendaciones(doc, ctx)
    _add_datos_responsable(doc, ctx)
    _add_anexos(doc, ctx)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ===========================================================================
# INFORME DIAGNÓSTICO (formato aprobado por dirección jul-2026)
#
# Las funciones *_inf reproducen el documento de referencia
# "Informe_Diagnostico_NOM035_San_Luis_2026 (2)-1.docx": misma redacción,
# tablas, negritas y estructura, con los datos de cada planta tomados del
# ctx. Las funciones del borrador (arriba) NO se tocan.
# ===========================================================================

_ALIGN_INF = {
    'CENTER':  WD_ALIGN_PARAGRAPH.CENTER,
    'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _p_segs(doc, segs, style=None, align=None):
    """Párrafo a partir de segmentos (texto, negrita). La negrita solo se
    escribe cuando aplica (los runs normales heredan del estilo)."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align:
        p.alignment = align
    for texto, bold in segs:
        run = p.add_run(texto)
        if bold:
            run.bold = True
    return p


def _emit_bloque(doc, bloque, listas=None):
    """Emite un bloque estático extraído del documento de referencia
    (contenido_informe.BLOQUE_*): párrafos con estilo, alineación y
    segmentos de negrita idénticos.

    Los ítems con clave 'lista' se emiten como listas REALES de Word
    (docx_geometry.nueva_lista/parrafo_lista): cada grupo distinto es una
    instancia de numeración nueva, de modo que cada sublista reinicia en 1
    — igual que el documento maestro de Zapotitlán."""
    listas = {} if listas is None else listas
    for item in bloque:
        align = _ALIGN_INF.get(item['align'])
        if item.get('lista'):
            tipo, grupo = item['lista']
            if grupo not in listas:
                listas[grupo] = geom.nueva_lista(doc, tipo)
            p = geom.parrafo_lista(doc, listas[grupo], ilvl=item.get('nivel', 0))
            if align:
                p.alignment = align
            for texto, bold in item['segs']:
                run = p.add_run(texto)
                if bold:
                    run.bold = True
            continue
        style = item['style'] if item['style'] != 'Normal' else None
        _p_segs(doc, item['segs'], style=style, align=align)


# --- Diseño de tabla aprobado por dirección (jul-2026, "opción 2"):
# encabezado rojo corporativo con texto blanco, filas alternadas en gris muy
# claro y bordes finos grises. Sustituye al estilo 'Grid Table 1 Light Accent
# 2' del maestro. Solo cambia la apariencia: contenido, anchos de columna y
# repetición de encabezado siguen viniendo de docx_geometry. ---
_TABLA_ENCABEZADO = 'C8102E'
_TABLA_BANDA = 'F2F2F2'
_TABLA_BORDE = 'D9D9D9'


# Semáforo de nivel de riesgo (misma paleta que las gráficas de barras del
# informe y que el frontend `--nom-riesgo-*`). La NOM-035 no define colores:
# es una convención interna, aprobada por dirección jul-2026.
NIVEL_COLOR_DOCX = {
    'nulo':     '10B981',
    'bajo':     '84CC16',
    'medio':    'F59E0B',
    'alto':     'EF4444',
    'muy_alto': '7C3AED',
}


def _texto_contrastante(hex6):
    """Negro o blanco, el que tenga mayor contraste sobre `hex6` (WCAG 2.1).
    Con esta paleta el ámbar y el lima exigen texto negro; el morado, blanco."""
    canales = [int(hex6[i:i + 2], 16) / 255 for i in (0, 2, 4)]
    lineal = [c / 12.92 if c <= 0.03928 else ((c + 0.055) / 1.055) ** 2.4 for c in canales]
    luminancia = 0.2126 * lineal[0] + 0.7152 * lineal[1] + 0.0722 * lineal[2]
    contraste_negro = (luminancia + 0.05) / 0.05
    contraste_blanco = 1.05 / (luminancia + 0.05)
    return RGBColor(0, 0, 0) if contraste_negro >= contraste_blanco else RGBColor(0xFF, 0xFF, 0xFF)


def _pintar_celda_nivel(cell, clave_nivel, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Colorea la celda con el semáforo del nivel. Reemplaza el sombreado que
    haya puesto el diseño de tabla (banding) en vez de acumular otro w:shd,
    que dejaría la celda con dos rellenos en el XML."""
    color_fondo = NIVEL_COLOR_DOCX.get(clave_nivel)
    if not color_fondo:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    _cell_bg(cell, color_fondo)
    _celda_texto(cell, cell.text, bold=True, color=_texto_contrastante(color_fondo),
                 align=align)


_ETIQUETA_A_NIVEL = {
    'nulo': 'nulo', 'nulo / despreciable': 'nulo', 'nulo o despreciable': 'nulo',
    'bajo': 'bajo', 'medio': 'medio', 'alto': 'alto', 'muy alto': 'muy_alto',
}


def _nivel_de_etiqueta(texto):
    """Clave de nivel a partir del texto de una celda ('Muy alto %', 'Nulo /
    Despreciable'…). Devuelve None para las columnas acumuladas ('Medio o
    superior %', 'Alto o Muy alto %'), que no corresponden a un solo nivel."""
    limpio = (texto or '').strip().rstrip('%').strip().lower()
    return _ETIQUETA_A_NIVEL.get(limpio)


def _pintar_encabezado_niveles(table):
    """Aplica el semáforo a las celdas del encabezado que nombran un nivel
    (tablas cuyos niveles van en columnas, no en filas)."""
    for cell in table.rows[0].cells:
        clave = _nivel_de_etiqueta(cell.text)
        if clave:
            _pintar_celda_nivel(cell, clave, align=WD_ALIGN_PARAGRAPH.CENTER)


def _fijar_bordes_tabla(table, exterior=8, interior=4, color=_TABLA_BORDE):
    """Reemplaza los bordes que hereda el estilo por una rejilla fina."""
    tblPr = table._tbl.tblPr
    for viejo in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(viejo)
    borders = OxmlElement('w:tblBorders')
    for lado, sz in (('top', exterior), ('left', exterior), ('bottom', exterior),
                     ('right', exterior), ('insideH', interior), ('insideV', interior)):
        el = OxmlElement(f'w:{lado}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), color)
        el.set(qn('w:space'), '0')
        borders.append(el)
    tblPr.append(borders)


def _celda_texto(cell, texto, bold=False, size=9, color=None, align=None):
    """Reescribe el contenido de la celda con el formato del diseño aprobado
    (9 pt, 2 pt de aire arriba/abajo)."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = align
    run = p.add_run(str(texto))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _aplicar_diseno_tabla(table):
    """Encabezado rojo + banding gris. Se llama DESPUÉS de fijar la geometría
    para que el sombreado se añada sobre el tcPr ya construido (tcW antes de
    shd, como exige el esquema OOXML)."""
    _fijar_bordes_tabla(table)
    for cell in table.rows[0].cells:
        _cell_bg(cell, _TABLA_ENCABEZADO)
        _celda_texto(cell, cell.text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                     align=WD_ALIGN_PARAGRAPH.CENTER)
    for idx, row in enumerate(table.rows[1:]):
        for j, cell in enumerate(row.cells):
            if idx % 2 == 1:
                _cell_bg(cell, _TABLA_BANDA)
            _celda_texto(cell, cell.text, bold=(j == 0),
                         align=WD_ALIGN_PARAGRAPH.LEFT if j == 0 else None)


def _tabla_grid(doc, headers, rows, espacio=True, familia=None, anchos=None):
    """Tabla de datos con el diseño aprobado (_aplicar_diseno_tabla).

    Con `familia`/`anchos` se aplica geometría explícita (tblW/tblGrid/tcW,
    layout fijo, encabezado repetido): ver docx_geometry.FAMILIAS_TABLA."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for row_vals in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row_vals):
            cells[i].text = str(v)
    if familia or anchos:
        geom.fijar_geometria_tabla(table, anchos_twips=anchos, familia=familia)
    _aplicar_diseno_tabla(table)
    if espacio:
        doc.add_paragraph()
    return table


def _add_portada_inf(doc, ctx):
    """Portada del formato aprobado: logo, título 36pt, NOM en rojo,
    'PLANTA:' + nombre, fecha de generación. Sin sello."""
    logo_path = finders.find('documents/img/lear_logo.png')
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path:
        p_logo.add_run().add_picture(logo_path, width=Cm(6))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    run = p_title.add_run('INFORME DIAGNÓSTICO')
    run.bold = True
    run.font.size = Pt(36)

    p_vacio = doc.add_paragraph()
    p_vacio.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_nom = doc.add_paragraph()
    p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_nom.add_run('NOM-035-STPS-2018')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xC8, 0x10, 0x2E)

    p_planta = doc.add_paragraph()
    p_planta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_planta.paragraph_format.space_before = Pt(16)
    run = p_planta.add_run('PLANTA:')
    run.bold = True
    run.font.size = Pt(14)

    p_tenant = doc.add_paragraph()
    p_tenant.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tenant.paragraph_format.space_before = Pt(16)
    run = p_tenant.add_run(str(ctx['tenant'].nombre))
    run.bold = True
    run.font.size = Pt(14)

    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_fecha.add_run(f"Generado el {ctx['fecha_generado']}")
    run.font.size = Pt(12)

    p_fin = doc.add_paragraph()
    p_fin.paragraph_format.space_before = Pt(12)
    doc.add_page_break()


def _add_indice_inf(doc):
    """Índice: campo TOC (Word lo puebla al abrir gracias a updateFields de
    la plantilla) con el estilo/tabulador del documento de referencia."""
    _heading(doc, '1. Índice', level=1)
    para = doc.add_paragraph(style='toc 1')
    para.paragraph_format.tab_stops.add_tab_stop(
        Twips(8630), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    # El resultado provisional es indispensable: sin él LibreOffice no
    # reconoce el campo como índice y la actualización determinista del TOC
    # (docx_postprocess) no tendría nada que poblar.
    _add_field(para, ' TOC \\o "1-3" \\h \\z \\u ',
               resultado='Índice pendiente de actualización.')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_page_break()


def _add_objetivo_inf(doc, ctx):
    """§3 del formato aprobado: los objetivos específicos van con viñeta
    ⚫ (lista real de Word), como en el maestro de Zapotitlán — no con
    números ni con símbolos escritos como texto."""
    _heading(doc, '3. Objetivo', level=1)
    _heading(doc, '3.1. Objetivo general', level=3)
    doc.add_paragraph(ctx['objetivo_general'])
    _heading(doc, '3.2. Objetivos específicos', level=3)
    num_id = geom.nueva_lista(doc, 'bullet_negra')
    for obj in ctx['objetivos_especificos']:
        p = geom.parrafo_lista(doc, num_id)
        p.add_run(obj)


def _add_definiciones_inf(doc, ctx):
    """§4 del formato aprobado: viñetas • (Symbol) con el término en
    negrita, como el maestro."""
    _heading(doc, '4. Definiciones', level=1)
    num_id = geom.nueva_lista(doc, 'bullet')
    for termino, definicion in ctx['definiciones']:
        p = geom.parrafo_lista(doc, num_id)
        p.add_run(f'{termino}: ').bold = True
        p.add_run(definicion)


def _add_justificacion_muestra_inf(doc, ctx):
    jm = ctx['justificacion_muestra']
    _heading(doc, '5. Justificación de la muestra', level=1)

    _p_segs(doc, inf.JUSTIFICACION_P1_SEGS)
    for parrafo in jm['parrafos'][1:]:
        doc.add_paragraph(parrafo)

    eq1 = jm.get('ecuacion1')
    if eq1:
        def _pct(parte, total):
            return f'{parte / total * 100:.1f}%' if total else '0.0%'

        table = doc.add_table(rows=3, cols=6)
        table.style = 'Table Grid'

        encabezados = ['Colaboradores', 'Hombres', 'Mujeres', 'Muestra', 'Hombres', 'Mujeres']
        for j, h in enumerate(encabezados):
            cell = table.rows[0].cells[j]
            _cell_bg(cell, 'C8102E')
            _cell_run(cell, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))

        subencabezados = ['Total', '(N / %)', '(N / %)', f"N mín. {eq1['n_min']}", '(N / %)', '(N / %)']
        for j, s in enumerate(subencabezados):
            cell = table.rows[1].cells[j]
            _cell_bg(cell, '1A1A2E')
            _cell_run(cell, s, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))

        datos = [
            str(eq1['total']),
            f"{eq1['hombres']} / {_pct(eq1['hombres'], eq1['total'])}",
            f"{eq1['mujeres']} / {_pct(eq1['mujeres'], eq1['total'])}",
            str(eq1['muestra']),
            f"{eq1['muestra_hombres']} / {_pct(eq1['muestra_hombres'], eq1['muestra'])}",
            f"{eq1['muestra_mujeres']} / {_pct(eq1['muestra_mujeres'], eq1['muestra'])}",
        ]
        for j, d in enumerate(datos):
            _cell_run(table.rows[2].cells[j], d, size=10)
        geom.fijar_geometria_tabla(table, familia='sextos_6')

        doc.add_paragraph()

    if ctx.get('flujo_muestra'):
        p_intro = doc.add_paragraph(
            'Flujo de la muestra (Guía III) — de la población total a los '
            'cuestionarios analizados:'
        )
        p_intro.runs[0].italic = True
        geom.unir_con_siguiente(p_intro)
        _tabla_grid(
            doc,
            ['Etapa', 'N'],
            [[f['etapa'], f['n']] for f in ctx['flujo_muestra']],
            familia='flujo_2',
        )
        doc.add_paragraph()


def _add_metodologia_inf(doc, ctx):
    met = ctx['metodologia']
    _heading(doc, '6. Metodología', level=1)
    _heading(doc, '6.1. Objetivo de la evaluación', level=3)
    doc.add_paragraph(met['objetivo_evaluacion'])

    _heading(doc, '6.2. Instrumentos utilizados', level=3)
    _tabla_grid(doc, ['Instrumento', 'Descripción'], met['instrumentos'],
                familia='descripcion_2')

    _heading(doc, '6.3. Procedimiento de aplicación', level=3)
    # 6.3 completo + 6.4 Procesamiento a digital + 6.5 Evaluación y análisis:
    # redacción aprobada por dirección, extraída textualmente del documento
    # de referencia.
    _emit_bloque(doc, inf.BLOQUE_METODOLOGIA_PROC)

    if ctx['tasas_respuesta']:
        _heading(doc, 'Tasa de respuesta por instrumento', level=3)
        _tabla_grid(
            doc,
            ['Instrumento', 'Aplicados', 'Respondidos', 'Tasa'],
            [[t['guia'], t['total'], t['completadas'], f"{t['pct']}%"] for t in ctx['tasas_respuesta']],
            familia='tasa_4',
        )


def _nota_denominador_muestra(tabla):
    """Nota fija tras cada tabla demográfica: dato faltante SEPARADO del
    denominador de las categorías válidas (nunca mezclados en la misma
    cifra)."""
    if tabla['n_faltante'] == 0:
        return (
            f"N válido: {tabla['n_valido']}. Sin datos faltantes. "
            f"Los porcentajes se calculan sobre el N válido."
        )
    pct_f = tabla['pct_faltante']
    pct_f_txt = f"{pct_f:.0f}" if pct_f == int(pct_f) else f"{pct_f:.1f}"
    return (
        f"N válido: {tabla['n_valido']}. Datos faltantes: {tabla['n_faltante']} "
        f"({pct_f_txt}%). Los porcentajes se calculan sobre el N válido."
    )


def _add_informe_demografico_inf(doc, ctx):
    _heading(doc, '7. Informe demográfico', level=1)
    doc.add_paragraph(
        f"El perfil demográfico corresponde a la población analítica del informe: "
        f"{ctx['muestra']['total']} trabajadores con cuestionario de la Guía III Y Guía V "
        f"válido. Cada variable reporta su N válido; los porcentajes se calculan sobre "
        f"ese denominador."
    )
    doc.add_paragraph(
        'Nota general: pequeñas diferencias frente al 100% en la suma de porcentajes '
        'de una tabla pueden deberse exclusivamente al redondeo visual (cada '
        'categoría se redondea de forma independiente).'
    ).runs[0].italic = True
    rd = ctx.get('report_data')
    if rd and rd['faltantes_criticos']['variables']:
        doc.add_paragraph(
            'Alerta: las variables siguientes superan el umbral de datos faltantes '
            f"({rd['faltantes_criticos']['umbral_pct']}%) y se presentan sin gráfica: "
            + '; '.join(f"{v['variable']} ({v['pct_faltante']}% sin dato)"
                        for v in rd['faltantes_criticos']['variables'])
        ).runs[0].italic = True
    graficas_muestra = ctx.get('graficas', {}).get('muestra', {})

    _heading(doc, '7.1. Informe de datos generales', level=3)
    cols_generales = {'Sexo', 'Grupo de edad', 'Estado civil', 'Nivel de estudios'}
    for tabla in ctx['muestra_tablas']:
        if tabla['col'] in cols_generales:
            _heading(doc, tabla['titulo'], level=4)
            _tabla_grid(
                doc, [tabla['col'], 'Trabajadores', '%'],
                [[d['label'], d['count'], f"{d['pct']}%"] for d in tabla['datos']],
                familia='tercios_3',
            )
            doc.add_paragraph(_nota_denominador_muestra(tabla)).runs[0].italic = True
            _add_grafica(doc, graficas_muestra.get(tabla['col']))

    _heading(doc, '7.1.1. Informe de datos laborales', level=3)
    for tabla in ctx['muestra_tablas']:
        if tabla['col'] not in cols_generales:
            _heading(doc, tabla['titulo'], level=4)
            _tabla_grid(
                doc, [tabla['col'], 'Trabajadores', '%'],
                [[d['label'], d['count'], f"{d['pct']}%"] for d in tabla['datos']],
                familia='tercios_3',
            )
            doc.add_paragraph(_nota_denominador_muestra(tabla)).runs[0].italic = True
            _add_grafica(doc, graficas_muestra.get(tabla['col']))


def _add_ats_inf(doc, ctx):
    _heading(doc, '8. Informe de acontecimientos traumáticos severos (ATS)', level=1)
    # Introducción normativa (numerales 4.1 y 5.5) + criterio de la Guía I,
    # redacción aprobada por dirección.
    _emit_bloque(doc, inf.BLOQUE_ATS_INTRO)

    rd = ctx.get('report_data')
    if not rd:
        return
    ats = rd['tablas']['ats']
    _tabla_grid(
        doc,
        ['Indicador', 'N', '%'],
        [[f['indicador'], f['n'], f"{f['pct']}%" if f['pct'] is not None else '—']
         for f in ats['filas']],
        familia='ats_resumen_3',
    )
    doc.add_paragraph(
        f"Denominador: {ats['denominador']} (N = {ats['n_valido']}). "
        'El desglose por reactivo se presenta en el Anexo técnico.'
    ).runs[0].italic = True


def _tabla_niveles_rd_inf(doc, tabla_rd, col1, sufijo_nota=''):
    """Tabla estándar de niveles con estilo Grid; la nota de denominador ya
    no incluye la frase 'No existe un nivel global oficial…' (retirada por
    dirección)."""
    filas = sorted(tabla_rd['filas'], key=lambda f: f['prioridad'])
    _pintar_encabezado_niveles(_tabla_grid(
        doc,
        [col1, 'N', 'Nulo %', 'Bajo %', 'Medio %', 'Alto %', 'Muy alto %',
         'Medio o superior %', 'Alto o Muy alto %', 'Prioridad'],
        [[f['nombre'], f['n']]
         + [_fmt_pct(f['pct_nivel'][k]) for k in ('nulo', 'bajo', 'medio', 'alto', 'muy_alto')]
         + [_fmt_pct(f['pct_medio_o_mas']), _fmt_pct(f['pct_alto_o_muy_alto']), f"{f['prioridad']}°"]
         for f in filas],
        familia='niveles_10',
    ))
    doc.add_paragraph(
        f"Denominador: {tabla_rd['denominador']}. {inf.NOTA_NIVELES}{sufijo_nota}"
    ).runs[0].italic = True


def _add_informe_diagnostico_inf(doc, ctx):
    _heading(doc, '9. Informe diagnóstico sobre los factores de riesgo psicosocial y del entorno organizacional', level=1)

    _heading(doc, '9.1. Calificación final', level=3)
    doc.add_paragraph(
        f"Nivel de riesgo predominante en la organización: {ctx['nivel_global_label']} "
        f"(moda de los niveles individuales; N = {sum(d['count'] for d in ctx['distribucion'])}). "
        f"El {ctx.get('pct_medio_o_mas', 0)}% de las personas evaluadas se ubica en nivel "
        f"Medio o superior y el {ctx.get('pct_alto_muy_alto', 0)}% en Alto o Muy alto. "
        f"Puntaje promedio: {ctx['promedio_global']}; mediana: {ctx.get('mediana_global', '—')}."
    )
    if ctx.get('nota_promedio'):
        doc.add_paragraph(ctx['nota_promedio']).runs[0].italic = True
    doc.add_paragraph(ctx['nivel_global_texto'])
    tabla_dist = _tabla_grid(
        doc,
        ['Nivel de riesgo', 'Trabajadores', '%'],
        [[d['label'], d['count'], f"{d['pct']}%"] for d in ctx['distribucion']],
        familia='tercios_3',
    )
    for fila, d in zip(tabla_dist.rows[1:], ctx['distribucion']):
        _pintar_celda_nivel(fila.cells[0], d['key'])
    _add_grafica(doc, ctx.get('graficas', {}).get('distribucion'), width_cm=11)

    if ctx.get('acciones'):
        p_t7 = doc.add_paragraph(
            'Tabla 7, Guía de Referencia III, NOM-035-STPS-2018 — Criterios para la toma '
            'de acciones (transcripción del texto oficial):'
        )
        p_t7.runs[0].italic = True
        geom.unir_con_siguiente(p_t7)
        tabla_t7 = _tabla_grid(
            doc,
            ['Nivel de riesgo', 'Acción requerida'],
            [[a['label'], a['accion']] for a in ctx['acciones']],
            familia='descripcion_2',
        )
        # Semáforo por nivel en la columna "Nivel de riesgo" (las filas van en
        # el mismo orden que ctx['acciones'], que ya trae la clave del nivel).
        for fila, accion in zip(tabla_t7.rows[1:], ctx['acciones']):
            _pintar_celda_nivel(fila.cells[0], accion['key'])

    graficas_categorias = ctx.get('graficas', {}).get('categorias', {})
    graficas_dominios = ctx.get('graficas', {}).get('dominios', {})
    _niveles = ('nulo', 'bajo', 'medio', 'alto', 'muy_alto')

    _heading(doc, '9.2. Calificaciones por categoría', level=3)
    if ctx.get('rangos_categoria'):
        p_rc = doc.add_paragraph(inf.INTRO_RANGOS_CATEGORIA)
        p_rc.runs[0].italic = True
        geom.unir_con_siguiente(p_rc)
        _pintar_encabezado_niveles(_tabla_grid(
            doc,
            ['Categoría'] + [NIVEL_LABEL_DOCX[n] for n in _niveles],
            [[r['nombre']] + [r['rangos'][n] for n in _niveles] for r in ctx['rangos_categoria']],
            familia='rangos_6',
        ))
    _tabla_niveles_rd_inf(doc, ctx['report_data']['tablas']['categorias'], 'Categoría',
                          sufijo_nota=' ')
    for c in ctx['categorias_agregadas']:
        _add_grafica(doc, graficas_categorias.get(c['nombre']), width_cm=8)

    _heading(doc, '9.3. Calificaciones por dominio', level=3)
    if ctx.get('rangos_dominio'):
        p_rd = doc.add_paragraph(inf.INTRO_RANGOS_DOMINIO)
        p_rd.runs[0].italic = True
        geom.unir_con_siguiente(p_rd)
        _pintar_encabezado_niveles(_tabla_grid(
            doc,
            ['Dominio'] + [NIVEL_LABEL_DOCX[n] for n in _niveles],
            [[r['nombre']] + [r['rangos'][n] for n in _niveles] for r in ctx['rangos_dominio']],
            familia='rangos_6',
        ))
    _tabla_niveles_rd_inf(doc, ctx['report_data']['tablas']['dominios'], 'Dominio')
    for d in ctx['dominios_oficiales_agregados']:
        _add_grafica(doc, graficas_dominios.get(d['nombre']), width_cm=8)
    doc.add_paragraph()

    rd = ctx['report_data']

    # ---- 9.5. Áreas prioritarias (la 9.4 de dimensiones pasó al Anexo
    # técnico A.T.3 en el formato aprobado; la numeración 9.5/9.6 se
    # conserva tal cual el documento de referencia) ----
    areas = rd['tablas']['areas']
    _heading(doc, '9.5. Áreas prioritarias', level=3)
    p_ar = doc.add_paragraph(f"{areas['nota']} Denominador: {areas['denominador']}.")
    p_ar.runs[0].italic = True
    geom.unir_con_siguiente(p_ar)
    if areas['filas']:
        _pintar_encabezado_niveles(_tabla_grid(
            doc,
            ['Área', 'N válido', 'Nulo %', 'Bajo %', 'Medio %', 'Alto %', 'Muy alto %',
             'Medio o superior %', 'Alto o Muy alto %', 'Prioridad'],
            [[a['nombre'], a['n']]
             + [_fmt_pct(a['pct_nivel'][k]) for k in ('nulo', 'bajo', 'medio', 'alto', 'muy_alto')]
             + [_fmt_pct(a['pct_medio_o_mas']), _fmt_pct(a['pct_alto_o_muy_alto']), f"{a['prioridad']}°"]
             for a in areas['filas']],
            familia='niveles_10',
        ))
    else:
        doc.add_paragraph('No hay áreas con N suficiente para el análisis desagregado.')
    if areas['reservadas']:
        doc.add_paragraph(
            'Áreas reservadas por confidencialidad (N menor al umbral de '
            f"{areas['umbral_confidencialidad']}): "
            + ', '.join(f"{a['nombre']} (N={a['n']})" for a in areas['reservadas'])
        ).runs[0].italic = True

    # ---- 9.6. Violencia laboral (resumen del dominio oficial) ----
    viol = rd['tablas']['violencia']
    _heading(doc, '9.6. Violencia laboral (resumen)', level=3)
    p_vi = doc.add_paragraph(viol['nota'])
    p_vi.runs[0].italic = True
    geom.unir_con_siguiente(p_vi)
    tabla_viol = _tabla_grid(
        doc,
        ['Nivel', 'N', '%'],
        [[f['label'], f['n'], _fmt_pct(f['pct'])] for f in viol['filas']],
        familia='tercios_3',
    )
    for fila, f in zip(tabla_viol.rows[1:], viol['filas']):
        _pintar_celda_nivel(fila.cells[0], f['nivel'])
    doc.add_paragraph(
        f"Denominador: {viol['denominador']} (N = {viol['n_valido']}). "
        f"Personal en Alto o Muy alto: {_fmt_pct(viol['pct_alto_o_muy_alto'])}."
    ).runs[0].italic = True


def _add_conclusiones_inf(doc, ctx):
    c = ctx['conclusiones']
    dist = {d['key']: d['count'] for d in ctx['distribucion']}
    n = sum(dist.values()) or 1
    moda_n = dist.get(c['nivel_predominante'], 0)
    nivel_label = NIVEL_LABEL_DOCX.get(c['nivel_predominante'], c['nivel_predominante'])

    _heading(doc, '10. Conclusiones', level=1)
    doc.add_paragraph(
        inf.CONCLUSIONES_INTRO_1.format(n=c['muestra']),
        style='pdq2pg_selectionanchorcontainer',
    )
    doc.add_paragraph(inf.CONCLUSIONES_INTRO_2, style='Normal (Web)')
    # Separador simple: un Heading vacío contaminaba el índice y la
    # secuencia canónica de encabezados.
    doc.add_paragraph()

    _heading(doc, '10.1. Calificación final', level=3)
    doc.add_paragraph()
    _p_segs(doc, [(t.format(n=c['muestra'], acc=c['accion_global'],
                            pct_acc=c['pct_accion']), b)
                  for t, b in inf.CONCLUSION_10_1_P1])
    _p_segs(doc, [(t.format(alto=c['riesgo_alto'], pct_alto=c['pct_riesgo']), b)
                  for t, b in inf.CONCLUSION_10_1_P2])
    _p_segs(doc, [(t.format(nivel=nivel_label, moda=moda_n,
                            pct_moda=round(moda_n / n * 100, 1)), b)
                  for t, b in inf.CONCLUSION_10_1_P3])
    doc.add_paragraph()
    tabla_10_1 = _tabla_grid(
        doc,
        ['Muestra', 'Requieren Programa de intervención (M+A+MA)',
         'Prioridad interna (Alto+Muy alto)', 'Nivel predominante'],
        [[c['muestra'], f"{c['accion_global']} ({c['pct_accion']}%)",
          f"{c['riesgo_alto']} ({c['pct_riesgo']}%)", nivel_label]],
        familia='conclusion_4',
    )
    _pintar_celda_nivel(tabla_10_1.rows[1].cells[3], c['nivel_predominante'],
                        align=WD_ALIGN_PARAGRAPH.CENTER)

    _heading(doc, '10.2. Análisis por categoría', level=3)
    geom.unir_con_siguiente(doc.add_paragraph(inf.CONCLUSION_10_2_INTRO))
    _add_conclusiones_tabla_rankeada(doc, 'Categoría', c['categorias_rankeadas'],
                                     familia='rankeada_5', diseno_informe=True)

    _heading(doc, '10.3. Análisis por dominio', level=3)
    geom.unir_con_siguiente(doc.add_paragraph(
        'Los dominios se presentan ordenados de mayor a menor porcentaje de población en '
        'niveles Alto y Muy alto, permitiendo identificar los aspectos específicos del entorno '
        'laboral que concentran mayor riesgo psicosocial.'
    ))
    _add_conclusiones_tabla_rankeada(doc, 'Dominio', c['dominios_rankeados'],
                                     familia='rankeada_5', diseno_informe=True)

    cat_top = c['categorias_rankeadas'][0]['nombre'] if c['categorias_rankeadas'] else '—'
    dom_top = c['dominios_rankeados'][0]['nombre'] if c['dominios_rankeados'] else '—'
    _p_segs(doc, [(t.format(pct_acc=c['pct_accion'], cat=cat_top, dom=dom_top), b)
                  for t, b in inf.SINTESIS_SEGS])


def _add_recomendaciones_inf(doc, ctx):
    _heading(doc, '11. Recomendaciones', level=1)

    _heading(doc, '11.1. Recomendaciones generales', level=3)
    _tabla_grid(
        doc, ['Acción', 'Descripción'],
        [[titulo, texto] for titulo, texto in ctx['acciones_generales']],
        familia='descripcion_2',
    )

    recs_extendidas = ctx.get('recomendaciones_extendidas') or []
    _heading(doc, '11.2. Recomendaciones específicas', level=3)
    if not recs_extendidas:
        doc.add_paragraph(
            'Ningún dominio alcanzó nivel de riesgo Alto o Muy alto a nivel agregado. '
            'Se recomienda mantener las condiciones organizacionales actuales y reaplicar '
            'la evaluación en el siguiente ciclo normativo.'
        )
        return
    for idx, rec in enumerate(recs_extendidas, 1):
        _heading(doc, f"{idx}. {rec['dominio']} ({rec['pct_intervencion']}%)", level=4)
        doc.add_paragraph(rec['intro'])
        num_id = geom.nueva_lista(doc, 'bullet')
        for titulo, texto in rec['bullets']:
            p = geom.parrafo_lista(doc, num_id)
            p.add_run(f'{titulo}: ').bold = True
            p.add_run(texto)


def _add_matriz_intervencion_inf(doc, ctx):
    """En el formato aprobado la matriz quedó como título + nota (la tabla
    operativa la llena el centro de trabajo). El encabezado forma parte de
    la estructura canónica (índice del maestro), así que se emite siempre
    que el informe lleve datos compuestos."""
    rd = ctx.get('report_data')
    if not rd:
        return
    matriz = rd['tablas']['matriz_intervencion']
    _heading(doc, 'Matriz de intervención', level=3)
    doc.add_paragraph(f"{matriz['nota']} {matriz['denominador']}.").runs[0].italic = True


def _add_datos_responsable_inf(doc, ctx):
    """Sección 12 del formato aprobado: responsables de la implementación y
    seguimiento (coordinación técnica, consultoría operativa y responsable
    interno del centro de trabajo)."""
    _heading(doc, '12. Datos del responsable de la evaluación', level=1)
    _emit_bloque(doc, inf.BLOQUE_RESPONSABLES)


def _add_anexo_tecnico_inf(doc, ctx):
    bloques = ctx.get('bloques_captura')
    desglose = ctx.get('ats_desglose')
    if not bloques and not desglose:
        return
    _heading(doc, 'Anexo técnico', level=1)

    if bloques:
        _heading(doc, 'A.T.1. Bloques internos de captura (análisis complementario no normativo)', level=3)
        p_bl = doc.add_paragraph(
            'Los bloques D1-D14 corresponden a la organización interna del cuestionario y '
            'no son unidades de calificación de la NOM-035; su semáforo es únicamente '
            'referencial.'
        )
        p_bl.runs[0].italic = True
        geom.unir_con_siguiente(p_bl)
        _tabla_grid(
            doc,
            ['Clave', 'Bloque', 'Dominio oficial', 'Promedio'],
            [[dim['clave'], dim['nombre'], dim['dominio_oficial'],
              f"{dim['puntaje_promedio']}/{dim['puntaje_max']}"]
             for dim in bloques],
            familia='bloques_4',
        )

    if desglose:
        _heading(doc, 'A.T.2. Guía I — desglose por reactivo', level=3)
        doc.add_paragraph(
            'Tablas exhaustivas por pregunta (frecuencias por sexo). Se presentan en '
            'anexo técnico por su nivel de detalle; el resumen oficial está en la '
            'sección 8 del informe.'
        ).runs[0].italic = True
        graficas_ats = ctx.get('graficas', {}).get('ats', {})
        for seccion in desglose['secciones']:
            _heading(doc, seccion['nombre'], level=4)
            n_seccion = seccion.get('n')
            if n_seccion is not None:
                nota_n = (
                    f"N = {n_seccion} (cuestionarios Guía I válidos)." if seccion['clave'] == 'D1'
                    else f"N = {n_seccion} (cuestionarios Guía I elegibles: Sección I con ≥1 \"Sí\")."
                )
                doc.add_paragraph(nota_n).runs[0].italic = True
            if not seccion['filas']:
                doc.add_paragraph('Sin preguntas registradas para esta sección.')
                continue
            filas = [
                {**f, 'texto': inf.ATS_ACENTOS.get(f['texto'], f['texto'])}
                for f in seccion['filas']
            ]
            _add_ats_tabla(doc, filas)
            if any(f['h_si'] or f['m_si'] for f in seccion['filas']):
                _add_grafica(doc, graficas_ats.get(seccion['clave']), width_cm=13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # A.T.3 — Dimensiones (antes §9.4; dirección la movió al anexo técnico)
    _heading(doc, 'A.T.3 Dimensiones (indicador descriptivo)', level=3)
    if ctx.get('nota_dimensiones'):
        doc.add_paragraph(ctx['nota_dimensiones']).runs[0].italic = True
    if ctx['dimensiones'] and 'pct_promedio' in ctx['dimensiones'][0]:
        _tabla_grid(
            doc,
            ['#', 'Dimensión', 'Dominio', 'N', '% promedio', '% mediana'],
            [[dim['orden'], dim['nombre'], dim['dominio_oficial'], dim['n'],
              f"{dim['pct_promedio']}%", f"{dim['pct_mediana']}%"]
             for dim in ctx['dimensiones']],
            espacio=False,
            familia='dimensiones_6',
        )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_informe_diagnostico_docx(ctx: dict) -> io.BytesIO:
    """Construye el Informe Diagnóstico completo con el formato aprobado por
    dirección (jul-2026): portada 'PLANTA:', índice TOC, panel ejecutivo,
    secciones 2-12 y Anexo técnico. Requiere `ctx` generado con
    `informe_extendido=True` (gráficas y recomendaciones extensas)."""
    doc = Document(str(_PLANTILLA_INFORME))

    _add_portada_inf(doc, ctx)
    _add_indice_inf(doc)
    _add_panel_ejecutivo(doc, ctx)
    _add_datos_centro_trabajo(doc, ctx)
    _add_objetivo_inf(doc, ctx)
    _add_definiciones_inf(doc, ctx)
    _add_justificacion_muestra_inf(doc, ctx)
    _add_metodologia_inf(doc, ctx)
    _add_informe_demografico_inf(doc, ctx)
    _add_ats_inf(doc, ctx)
    _add_informe_diagnostico_inf(doc, ctx)
    _add_conclusiones_inf(doc, ctx)
    _add_recomendaciones_inf(doc, ctx)
    _add_matriz_intervencion_inf(doc, ctx)
    _add_datos_responsable_inf(doc, ctx)
    _add_anexo_tecnico_inf(doc, ctx)
    _add_anexos(doc, ctx, con_geometria=True)

    # Control global de saltos: títulos con keep_with_next y viudas/huérfanas.
    geom.aplicar_control_saltos(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
