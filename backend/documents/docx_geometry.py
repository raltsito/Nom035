"""Geometría DOCX explícita del Informe Diagnóstico NOM-035.

El documento de referencia (PLANTA Zapot-2.pdf ≡ zapotitlan_informe_diagnostico_2026.pdf)
se maqueta en tamaño carta (12240×15840 twips) con márgenes izquierdo/derecho de
1800 twips (plantilla_informe.docx) → ancho útil 8640 twips. Aquí se definen:

- Anchos de columna por FAMILIA de tabla (proporciones del ancho útil), con
  `tblW`/`tblGrid`/`tcW` explícitos y layout fijo: ningún renderizador (Word,
  LibreOffice) debe re-balancear columnas ni partir palabras carácter por
  carácter por falta de ancho.
- Repetición de la fila de encabezado (`tblHeader`) y control de división de
  filas (`cantSplit`).
- `keep_with_next` para títulos y párrafos introductorios pegados a su tabla.
- Listas REALES de Word: cada lista lógica recibe su propia instancia de
  numeración (reinicia en 1; sin números escritos como texto ni símbolos
  Unicode sueltos). El maestro usa dos glifos de viñeta: • (Symbol) y ⚫
  (círculo negro medio).
"""
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Carta 12240 − márgenes 1800 + 1800 (sectPr de plantilla_informe.docx)
ANCHO_UTIL_TWIPS = 8640

# ---------------------------------------------------------------------------
# Familias de tabla → proporciones de columna (suman 1.0). Derivadas de la
# maqueta del documento maestro de Zapotitlán.
# ---------------------------------------------------------------------------
FAMILIAS_TABLA = {
    # Distribuciones demográficas / distribución final / violencia (etiqueta, N, %)
    'tercios_3':      (1 / 3, 1 / 3, 1 / 3),
    # Tablas de niveles por categoría/dominio/área (nombre + N + 5 niveles +
    # 2 acumulados + prioridad) — col. 1 ancha como en el maestro y columna
    # final suficiente para que "Prioridad" no se parta por carácter.
    'niveles_10':     (0.18, 0.05, 0.09, 0.09, 0.09, 0.09, 0.09, 0.103, 0.103, 0.114),
    # Rangos oficiales de corte (nombre + 5 niveles)
    'rangos_6':       (0.22, 0.156, 0.156, 0.156, 0.156, 0.156),
    # Instrumentos (6.2) / Tabla 7 / Recomendaciones generales (concepto, descripción)
    'descripcion_2':  (0.42, 0.58),
    # Flujo de la muestra (etapa, N)
    'flujo_2':        (0.75, 0.25),
    # Ecuación 1 (6 columnas iguales)
    'sextos_6':       (1 / 6,) * 6,
    # Tasa de respuesta por instrumento
    'tasa_4':         (0.34, 0.22, 0.22, 0.22),
    # Resumen ATS (indicador, N, %)
    'ats_resumen_3':  (0.50, 0.25, 0.25),
    # Conclusión 10.1 (muestra, M+A+MA, A+MA, nivel predominante)
    'conclusion_4':   (0.16, 0.32, 0.28, 0.24),
    # Conclusiones 10.2/10.3 rankeadas (nombre, nivel, % programa, % alto, prioridad)
    'rankeada_5':     (0.28, 0.20, 0.24, 0.16, 0.12),
    # Guía I desglose por reactivo (#, pregunta, H sí/no, M sí/no, total sí/no)
    'ats_desglose_8': (0.06, 0.40, 0.09, 0.09, 0.09, 0.09, 0.09, 0.09),
    # A.T.3 dimensiones (#, dimensión, dominio, N, % promedio, % mediana)
    'dimensiones_6':  (0.06, 0.26, 0.26, 0.10, 0.16, 0.16),
    # A.T.1 bloques internos (clave, bloque, dominio oficial, promedio)
    'bloques_4':      (0.10, 0.30, 0.36, 0.24),
    # Panel ejecutivo (3 tarjetas por fila)
    'panel_3':        (1 / 3, 1 / 3, 1 / 3),
    # Anexo 13.4 agrupación (categoría, dominio, claves)
    'agrupacion_3':   (0.30, 0.40, 0.30),
}

# Márgenes de celda reducidos para las familias más densas (twips).
FAMILIAS_MARGEN = {
    'niveles_10': 28,
}


def anchos_familia(familia):
    """Anchos en twips de una familia registrada; la suma es exactamente
    ANCHO_UTIL_TWIPS (el residuo del redondeo se asigna a la última columna)."""
    return anchos_por_proporcion(FAMILIAS_TABLA[familia])


def anchos_por_proporcion(proporciones):
    total = sum(proporciones)
    anchos = [int(round(p / total * ANCHO_UTIL_TWIPS)) for p in proporciones]
    anchos[-1] += ANCHO_UTIL_TWIPS - sum(anchos)
    return anchos


def anchos_equitativos(n_cols, primera=None):
    """N columnas iguales; si `primera` (proporción 0-1) se indica, la primera
    columna la usa y el resto se reparte en partes iguales."""
    if primera is None:
        return anchos_por_proporcion((1.0 / n_cols,) * n_cols)
    resto = (1.0 - primera) / (n_cols - 1)
    return anchos_por_proporcion((primera,) + (resto,) * (n_cols - 1))


# ---------------------------------------------------------------------------
# Aplicación de geometría a una tabla python-docx
# ---------------------------------------------------------------------------

def _set_child(parent, tag, attrs):
    """Reemplaza (o crea) el hijo `tag` de `parent` con los atributos dados."""
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    for k, v in attrs.items():
        el.set(qn(k), str(v))
    return el


def fijar_geometria_tabla(table, anchos_twips=None, familia=None,
                          repetir_encabezado=True, dividir_filas=True,
                          tamano_pt=None, margen_celda_twips=57):
    """Geometría explícita: `tblW` dxa = suma de `tblGrid`; `tcW` por celda;
    layout FIJO (sin autofit); márgenes de celda; encabezado repetido en cada
    página; y opcionalmente filas indivisibles.

    La suma de anchos siempre es el ancho útil de la página: ninguna tabla
    puede desbordar el área de texto.
    """
    if anchos_twips is None:
        anchos_twips = anchos_familia(familia) if familia else \
            anchos_equitativos(len(table.columns))
    if familia in FAMILIAS_MARGEN:
        margen_celda_twips = FAMILIAS_MARGEN[familia]
    n_cols = len(table.columns)
    if len(anchos_twips) != n_cols:
        raise ValueError(f'{n_cols} columnas pero {len(anchos_twips)} anchos')
    total = sum(anchos_twips)

    tbl = table._tbl
    tblPr = tbl.tblPr
    _set_child(tblPr, 'w:tblW', {'w:w': total, 'w:type': 'dxa'})
    _set_child(tblPr, 'w:tblLayout', {'w:type': 'fixed'})

    # Márgenes de celda uniformes (izq/der) como el documento de referencia.
    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is None:
        tblCellMar = OxmlElement('w:tblCellMar')
        tblPr.append(tblCellMar)
    for lado in ('w:left', 'w:right'):
        _set_child(tblCellMar, lado, {'w:w': margen_celda_twips, 'w:type': 'dxa'})

    # tblGrid exacto
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    for w in anchos_twips:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)

    for i_fila, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        if repetir_encabezado and i_fila == 0:
            _set_child(trPr, 'w:tblHeader', {'w:val': 'true'})
        if not dividir_filas:
            _set_child(trPr, 'w:cantSplit', {'w:val': 'true'})
        for celda, w in zip(row.cells, anchos_twips):
            tcPr = celda._tc.get_or_add_tcPr()
            _set_child(tcPr, 'w:tcW', {'w:w': w, 'w:type': 'dxa'})
            if tamano_pt:
                for p in celda.paragraphs:
                    for r in p.runs:
                        if r.font.size is None:
                            r.font.size = Pt(tamano_pt)
    return table


def unir_con_siguiente(paragraph):
    """keep_with_next: el párrafo no queda huérfano al final de la página
    (títulos y notas introductorias pegados a su tabla/primer párrafo)."""
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def aplicar_control_saltos(doc):
    """Pasada final sobre el documento: keep_with_next en todos los títulos
    (Heading 1-4 y sus alias en español) y control de viudas/huérfanas en el
    estilo Normal."""
    try:
        normal = doc.styles['Normal']
        normal.paragraph_format.widow_control = True
    except KeyError:
        pass
    encabezados = {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
                   'Título 1', 'Título 2', 'Título 3', 'Título 4'}
    for p in doc.paragraphs:
        if p.style is not None and p.style.name in encabezados:
            p.paragraph_format.keep_with_next = True
    return doc


# ---------------------------------------------------------------------------
# Listas reales de Word (numbering.xml)
# ---------------------------------------------------------------------------

# Glifos del documento maestro: • = Symbol F0B7; ⚫ = círculo negro medio.
_BULLET_SYMBOL = ('', 'Symbol')
_BULLET_NEGRA = ('⚫', 'Segoe UI Symbol')
_BULLET_SUB = ('o', 'Courier New')


def _lvl(ilvl, fmt, lvl_text, font=None, left=720, hanging=360):
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), str(ilvl))
    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    lvl.append(start)
    numFmt = OxmlElement('w:numFmt')
    numFmt.set(qn('w:val'), fmt)
    lvl.append(numFmt)
    lvlText = OxmlElement('w:lvlText')
    lvlText.set(qn('w:val'), lvl_text)
    lvl.append(lvlText)
    lvlJc = OxmlElement('w:lvlJc')
    lvlJc.set(qn('w:val'), 'left')
    lvl.append(lvlJc)
    pPr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left))
    ind.set(qn('w:hanging'), str(hanging))
    pPr.append(ind)
    lvl.append(pPr)
    if font:
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rFonts.set(qn(attr), font)
        rPr.append(rFonts)
        lvl.append(rPr)
    return lvl


def _elemento_numbering(doc):
    """Raíz <w:numbering> del documento (la plantilla del informe siempre
    trae word/numbering.xml)."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    part = doc.part.part_related_by(RT.NUMBERING)
    return getattr(part, 'element', None) if hasattr(part, 'element') else part._element


def nueva_lista(doc, tipo='decimal'):
    """Crea una instancia de numeración NUEVA (la lista reinicia en 1) y
    devuelve su numId. Tipos:

    - 'decimal': 1. 2. 3.  (pasos de procedimiento)
    - 'bullet': • (Symbol) con subnivel 'o' (criterios, responsables)
    - 'bullet_negra': ⚫ (objetivos específicos, "Durante", checklist 6.4)
    """
    numbering = _elemento_numbering(doc)

    abstract_ids = [int(a.get(qn('w:abstractNumId')))
                    for a in numbering.findall(qn('w:abstractNum'))]
    num_ids = [int(n.get(qn('w:numId'))) for n in numbering.findall(qn('w:num'))]
    abs_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstractNum = OxmlElement('w:abstractNum')
    abstractNum.set(qn('w:abstractNumId'), str(abs_id))
    multi = OxmlElement('w:multiLevelType')
    multi.set(qn('w:val'), 'hybridMultilevel')
    abstractNum.append(multi)

    if tipo == 'decimal':
        abstractNum.append(_lvl(0, 'decimal', '%1.'))
        abstractNum.append(_lvl(1, 'lowerLetter', '%2)', left=1440))
    elif tipo == 'bullet':
        abstractNum.append(_lvl(0, 'bullet', *_BULLET_SYMBOL))
        abstractNum.append(_lvl(1, 'bullet', *_BULLET_SUB, left=1440))
    elif tipo == 'bullet_negra':
        abstractNum.append(_lvl(0, 'bullet', *_BULLET_NEGRA))
        abstractNum.append(_lvl(1, 'bullet', *_BULLET_SUB, left=1440))
    else:
        raise ValueError(f'tipo de lista desconocido: {tipo}')

    # Esquema OOXML: todos los abstractNum preceden a los num.
    nums = numbering.findall(qn('w:num'))
    if nums:
        nums[0].addprevious(abstractNum)
    else:
        numbering.append(abstractNum)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abstractNumId = OxmlElement('w:abstractNumId')
    abstractNumId.set(qn('w:val'), str(abs_id))
    num.append(abstractNumId)
    numbering.append(num)
    return num_id


def parrafo_lista(doc, num_id, ilvl=0, style=None):
    """Párrafo vinculado a la instancia de numeración `num_id` (lista real de
    Word). El texto se agrega por el llamador con add_run."""
    p = doc.add_paragraph(style=style)
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_el)
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(numId_el)
    pPr.append(numPr)
    return p
