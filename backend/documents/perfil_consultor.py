"""Perfil profesional del consultor, al cierre del Informe Diagnóstico.

Va ÚNICAMENTE en el informe que entrega el botón "Descargar informe" de
/resultados. Los demás entregables (reporte psicológico, anexo fotográfico,
actas, difusión, plan de acción y política) NO lo llevan: es una decisión
deliberada, no un olvido.

El bloque se emite como texto con formato directo, NO como encabezado de
Word: la secuencia de encabezados del Informe Diagnóstico está fijada en
`estructura_canonica.py` y corresponde al maestro aprobado por dirección.
"""
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

TITULO = 'Perfil profesional del consultor'

NOMBRE = 'Lic. Javier Enrique Martínez Becerra'
CARGO = ('Consultor especializado en NOM-035-STPS-2018, salud mental y '
         'bienestar organizacional')
CEDULA = '0518000588'

PARRAFOS = (
    'Psicólogo y consultor organizacional con experiencia en la aplicación, '
    'calificación, análisis e interpretación de las Guías de Referencia de la '
    'NOM-035-STPS-2018, así como en la elaboración de diagnósticos técnicos, '
    'informes ejecutivos, programas de intervención y estrategias para la '
    'prevención de factores de riesgo psicosocial, violencia laboral y '
    'promoción de entornos organizacionales favorables.',

    'Ha participado en procesos de evaluación, capacitación y consultoría en '
    'más de 36 centros de trabajo, con un alcance aproximado de 34,000 '
    'trabajadores. Su experiencia comprende la validación de muestras y bases '
    'de datos, el análisis de categorías, dominios y áreas prioritarias, la '
    'identificación de acontecimientos traumáticos severos, el diseño de '
    'programas de bienestar laboral y la capacitación de directivos, '
    'supervisores, recursos humanos y comités internos.',

    'Cuenta con formación en psicología, evaluación y rehabilitación '
    'neuropsicológica, terapia cognitivo-conductual, terapias contextuales, '
    'salud mental y desarrollo humano. Es fundador y presidente del Instituto '
    'de Atención Integral y Desarrollo Humano A.C., fundador de Academia INTRA '
    'y participante en proyectos de consultoría organizacional orientados a '
    'convertir los resultados de la NOM-035 en acciones concretas, medibles y '
    'sostenibles.',
)

_GRIS = RGBColor(0x53, 0x5B, 0x66)


def agregar_a_docx(doc):
    """Añade el bloque al final del documento `doc` (python-docx).

    Usa formato directo en lugar de estilos con nombre para no depender de
    los estilos que traiga la plantilla base.
    """
    separador = doc.add_paragraph()
    separador.paragraph_format.space_before = Pt(18)
    separador.paragraph_format.space_after = Pt(0)

    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_after = Pt(6)
    p_titulo.paragraph_format.keep_with_next = True
    r_titulo = p_titulo.add_run(TITULO.upper())
    r_titulo.bold = True
    r_titulo.font.size = Pt(10)
    r_titulo.font.color.rgb = _GRIS

    p_nombre = doc.add_paragraph()
    p_nombre.paragraph_format.space_after = Pt(0)
    p_nombre.paragraph_format.keep_with_next = True
    r_nombre = p_nombre.add_run(NOMBRE)
    r_nombre.bold = True
    r_nombre.font.size = Pt(11)

    p_cargo = doc.add_paragraph()
    p_cargo.paragraph_format.space_after = Pt(8)
    p_cargo.paragraph_format.keep_with_next = True
    r_cargo = p_cargo.add_run(CARGO)
    r_cargo.font.size = Pt(9)
    r_cargo.font.color.rgb = _GRIS
    r_cedula = p_cargo.add_run(f'\nCédula profesional: {CEDULA}')
    r_cedula.font.size = Pt(9)
    r_cedula.font.color.rgb = _GRIS

    for i, texto in enumerate(PARRAFOS):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        # Los dos primeros párrafos no deben quedar solos al pie de página.
        p.paragraph_format.keep_with_next = i < len(PARRAFOS) - 1
        run = p.add_run(texto)
        run.font.size = Pt(9.5)
