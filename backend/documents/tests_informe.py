"""Pruebas de fidelidad del Informe Diagnóstico NOM-035 (formato aprobado).

Cubren los requisitos de la corrección integral de fidelidad (jul-2026):
- Flujo único: ninguna planta se sirve desde archivos fijos por nombre.
- Prueba integral con el fixture anonimizado de la planta PRUEBA
  (mismo endpoint y mismo motor que producción).
- Secuencia canónica exacta de encabezados (9.3 → 9.5 → 9.6; A.T.1-A.T.3).
- Snapshot del texto fijo aprobado (maestro de Zapotitlán).
- Auditoría OOXML de geometría de tablas (tblW/tblGrid/tcW coherentes).
- Índice REAL actualizado (LibreOffice) sin marcadores rotos.
- Validaciones bloqueantes de completitud (5 categorías / 10 dominios / 25
  dimensiones), coherencia de N y confidencialidad.
- Renderizado sin palabras partidas carácter por carácter ni páginas vacías.
"""
import io
import os
import re
import subprocess
import tempfile
import unittest
import zipfile
from pathlib import Path

from django.core.management import call_command
from django.test import SimpleTestCase, TestCase, override_settings
from docx import Document
from rest_framework.test import APIClient

from m06_results.scoring import _CATEGORIAS_OFICIALES, _DOMINIOS_OFICIALES
from .docx_postprocess import localizar_soffice
from .estructura_canonica import ESTRUCTURA_CANONICA, TEXTOS_CANONICOS
from .report_data import componer_report_data
from .tests import _cat, _dom, _persona, _raw
from .views import RAZON_SOCIAL_CORPORATIVA

MIME_DOCX = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

HAY_LIBREOFFICE = localizar_soffice() is not None

# Frases FIJAS del maestro (PLANTA Zapot-2.pdf) que todo informe debe
# reproducir literalmente, independientemente de la planta.
TEXTO_FIJO_MAESTRO = [
    # §3 objetivo general (inicio)
    'El objetivo de este informe es identificar, analizar y prevenir los factores de riesgo',
    # §5 justificación
    'utilizando un nivel de confianza del 95% (Z = 1.96), probabilidad de ocurrencia p = q = 0.5',
    'La selección de la muestra no fue aleatoria simple, sino estratificada',
    # §6.3
    'Se determinó el número mínimo de trabajadores a quienes se aplicarían los cuestionarios',
    'Se evitó conducir, persuadir o dirigir las respuestas.',
    'Al concluir los instrumentos, la plataforma realizó una verificación automática de',
    # §6.4
    'no se empleó un cuestionario alternativo sujeto al procedimiento de validación señalado en el numeral 7.5.',
    'Los registros parciales, inconsistentes o inválidos se conservaron para trazabilidad',
    'Los reactivos 65–68 sólo se incluyeron cuando el trabajador indicó que atendía clientes',
    # §6.5
    'No se asignó un nivel oficial único al centro de trabajo mediante el promedio, la mediana o',
    # §8 ATS
    'un acontecimiento traumático severo es aquel experimentado durante o con motivo del trabajo',
    'Los criterios de las secciones II, III y IV se evalúan de manera independiente',
    # §9.1 — Tabla 7 (transcripción oficial)
    'El riesgo resulta despreciable por lo que no se requiere medidas adicionales.',
    'mediante un Programa de intervención que deberá incluir evaluaciones específicas',
    # §9.2 / 9.3 (intros de rangos)
    'Rangos oficiales de calificación por categoría — Guía de Referencia III.',
    'Rangos oficiales de calificación por dominio — Guía de Referencia III.',
    # §10
    'se presentan las siguientes conclusiones.',
    'la cual contempla la implementación o fortalecimiento de un Programa de intervención',
    'Este resultado corresponde a la moda de la distribución de los cuestionarios individuales',
    # §11.1 (catálogo fijo)
    'Formular, documentar y difundir una política empresarial que declare el compromiso del',
    'Realizar una nueva aplicación de la Guía de Referencia III en un plazo no mayor a 12',
    # §12 responsables (redacción del maestro)
    'La responsabilidad del cumplimiento de la NOM-035-STPS-2018 corresponde al patrón',
    'Carlos Alberto González Becerra',
    'Canalizar a las personas que cumplan los criterios de la Guía I para ATS y activar los',
    # Anexo técnico
    'Los bloques D1-D14 corresponden a la organización interna del cuestionario',
    'Indicador descriptivo. La NOM-035-STPS-2018 no establece puntos de corte por dimensión.',
]

# Datos de otras plantas que JAMÁS deben aparecer en el informe de PRUEBA.
# La razón social corporativa NO entra en esta lista: todas las plantas son la
# misma persona moral, así que §2.1 la lleva fija (views.RAZON_SOCIAL_CORPORATIVA).
DATOS_OTRAS_PLANTAS = [
    'San Luis Potosí', 'Circuito Exportación', 'Zona Industrial',
    'Zapotitlán', 'Zapotitlan', 'Tláhuac',
]


def _texto_docx(doc):
    partes = [p.text for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                partes.append(celda.text)
    return '\n'.join(partes)


def _encabezados(doc, niveles=(1, 2, 3)):
    resultado = []
    for p in doc.paragraphs:
        nombre = p.style.name if p.style is not None else ''
        m = re.match(r'(?:Heading|Título|heading) (\d)$', nombre)
        if m and int(m.group(1)) in niveles and p.text.strip():
            resultado.append((p.text.strip(), int(m.group(1))))
    return resultado


class FixturePruebaMixin:
    """Carga el fixture PRUEBA (una vez por clase) y genera el informe por
    el MISMO endpoint de producción."""

    @classmethod
    def setUpTestData(cls):
        os.environ['NOM035_PERMITIR_FIXTURE'] = '1'
        call_command('cargar_fixture_prueba')
        from tenants.models import Tenant
        from m00_onboarding.models import CicloNOM
        cls.tenant = Tenant.objects.get(rfc='PRU010101PRB')
        cls.ciclo = CicloNOM.objects.get(tenant=cls.tenant)

    def _cliente_admin(self):
        from accounts.models import User
        user, _ = User.objects.get_or_create(
            username='admin_prueba',
            defaults={'email': 'admin_prueba@example.com', 'rol': 'tenant_admin',
                      'tenant': self.tenant},
        )
        cliente = APIClient()
        cliente.force_authenticate(user=user)
        return cliente

    def _descargar_informe(self, **params):
        cliente = self._cliente_admin()
        return cliente.get('/api/v1/documentos/informe-diagnostico/',
                           {'ciclo_id': self.ciclo.id, **params})


@override_settings(INFORME_TOC_ESTRICTO=False)
class TestInformeIntegralPrueba(FixturePruebaMixin, TestCase):
    """Prueba integral: fixture anonimizado → endpoint de producción → DOCX."""

    _docx_cache = None

    def _docx(self):
        if type(self)._docx_cache is None:
            res = self._descargar_informe()
            assert res.status_code == 200, getattr(res, 'content', b'')[:500]
            type(self)._docx_cache = res
        return type(self)._docx_cache

    def test_respuesta_docx_mime_y_nombre(self):
        res = self._docx()
        self.assertEqual(res['Content-Type'], MIME_DOCX)
        cd = res['Content-Disposition']
        self.assertIn('Informe_Diagnostico_NOM035_PRUEBA_2026.docx', cd)
        self.assertTrue(res.content.startswith(b'PK'))  # ZIP/OOXML real

    def test_secuencia_canonica_de_encabezados(self):
        doc = Document(io.BytesIO(self._docx().content))
        obtenidos = _encabezados(doc, niveles=(1, 2, 3))
        self.assertEqual(obtenidos, ESTRUCTURA_CANONICA)

    def test_texto_fijo_snapshot_maestro(self):
        texto = re.sub(r'\s+', ' ', _texto_docx(Document(io.BytesIO(self._docx().content))))
        faltantes = [f for f in TEXTO_FIJO_MAESTRO
                     if re.sub(r'\s+', ' ', f) not in texto]
        self.assertEqual(faltantes, [], f'Texto fijo ausente: {faltantes}')

    def test_sin_datos_de_otras_plantas(self):
        texto = _texto_docx(Document(io.BytesIO(self._docx().content)))
        presentes = [d for d in DATOS_OTRAS_PLANTAS if d in texto]
        self.assertEqual(presentes, [], f'Datos ajenos presentes: {presentes}')
        self.assertIn('PRUEBA', texto)
        # §2.1: razón social única del corporativo, igual en todas las plantas.
        self.assertIn(RAZON_SOCIAL_CORPORATIVA, texto)

    def test_auditoria_ooxml_anchos_de_tablas(self):
        """tblW dxa == suma de tblGrid == suma de tcW por fila, layout fijo y
        encabezado repetido en todas las tablas de datos."""
        from documents.docx_geometry import ANCHO_UTIL_TWIPS
        with zipfile.ZipFile(io.BytesIO(self._docx().content)) as z:
            xml = z.read('word/document.xml').decode('utf-8')
        tablas = re.findall(r'<w:tbl>.*?</w:tbl>', xml, re.S)
        self.assertGreater(len(tablas), 20)
        for i, t in enumerate(tablas):
            m_w = re.search(r'<w:tblW [^>]*w:w="(\d+)"[^>]*w:type="dxa"', t) or \
                  re.search(r'<w:tblW [^>]*w:type="dxa"[^>]*w:w="(\d+)"', t)
            self.assertIsNotNone(m_w, f'tabla {i} sin tblW dxa')
            total = int(m_w.group(1))
            self.assertEqual(total, ANCHO_UTIL_TWIPS,
                             f'tabla {i}: tblW {total} != ancho útil')
            grid = [int(w) for w in re.findall(r'<w:gridCol w:w="(\d+)"', t)]
            self.assertEqual(sum(grid), total, f'tabla {i}: tblGrid no suma tblW')
            self.assertIn('<w:tblLayout w:type="fixed"/>', t, f'tabla {i} sin layout fijo')
            for j, fila in enumerate(re.findall(r'<w:tr[ >].*?</w:tr>', t, re.S)):
                tcws = [int(w) for w in re.findall(
                    r'<w:tcW [^>]*w:w="(\d+)"', fila)]
                if tcws:
                    self.assertEqual(sum(tcws), total,
                                     f'tabla {i} fila {j}: tcW no suma tblW')

    def test_listas_reales_de_word(self):
        """Las listas del informe son numeración real (numPr), no texto."""
        with zipfile.ZipFile(io.BytesIO(self._docx().content)) as z:
            xml = z.read('word/document.xml').decode('utf-8')
            numbering = z.read('word/numbering.xml').decode('utf-8')
        self.assertGreater(xml.count('<w:numPr>'), 40)
        # Instancias nuevas creadas por docx_geometry (reinicio garantizado)
        self.assertGreater(numbering.count('<w:abstractNum '), 20)
        # Nadie escribe viñetas como carácter de texto
        doc = Document(io.BytesIO(self._docx().content))
        for p in doc.paragraphs:
            self.assertFalse(p.text.strip().startswith(('•', '⚫', '- ')),
                             f'Viñeta escrita como texto: {p.text[:60]!r}')

    def test_cifras_coherentes_con_report_data(self):
        """Las cifras del DOCX provienen de ReportDataNOM035 (fuente única)."""
        from documents.report_data import build_report_data
        rd = build_report_data(self.tenant, self.ciclo)
        self.assertTrue(rd['validaciones']['puede_emitirse'],
                        rd['validaciones']['errores_criticos'])
        texto = _texto_docx(Document(io.BytesIO(self._docx().content)))
        dist = rd['tablas']['distribucion_final']
        n = dist['n_valido']
        self.assertIn(f'N = {n}', texto)
        suma = sum(f['n'] for f in dist['filas'])
        self.assertEqual(suma, n)
        # Página del flujo de muestra: cuestionarios analizados == N válido
        self.assertIn('Cuestionarios analizados', texto)

    def test_confidencialidad_grupos_pequenos(self):
        from documents.report_data import build_report_data
        rd = build_report_data(self.tenant, self.ciclo)
        umbral = rd['meta']['umbral_confidencialidad']
        for area in rd['tablas']['areas']['filas']:
            self.assertGreaterEqual(area['n'], umbral)
        texto = _texto_docx(Document(io.BytesIO(self._docx().content)))
        for reservada in rd['tablas']['areas']['reservadas']:
            # El área reservada solo puede aparecer en la lista de reservadas,
            # nunca como fila del ranking (que llevaría porcentajes).
            self.assertNotRegex(
                texto, re.escape(reservada['nombre']) + r'\s+\d+\s+\d+(\.\d+)?%')

    def test_graficas_consumen_fuente_unica(self):
        """chart_pie de la distribución final recibe exactamente los conteos
        de ReportDataNOM035 (no recalcula)."""
        from unittest import mock
        from documents import views as dviews
        from documents.report_data import build_report_data

        rd = build_report_data(self.tenant, self.ciclo)
        esperados = {f['label']: f['n']
                     for f in rd['tablas']['distribucion_final']['filas'] if f['n']}

        capturas = {}
        original = dviews.graf.chart_pie

        def _espia(labels, counts, *a, **kw):
            if 'distribucion' not in capturas and set(labels) == set(esperados):
                capturas['distribucion'] = dict(zip(labels, counts))
            return original(labels, counts, *a, **kw)

        resultados_qs = self._resultados_qs()
        with mock.patch.object(dviews.graf, 'chart_pie', side_effect=_espia):
            dviews._build_psico_context(
                self.tenant, self.ciclo, resultados_qs, anonimo=False,
                informe_extendido=True, incluir_anexo_confidencial=False)
        self.assertEqual(capturas.get('distribucion'), esperados)

    def _resultados_qs(self):
        from m06_results.models import ResultadoAplicacion
        return ResultadoAplicacion.objects.filter(
            aplicacion__tenant=self.tenant, aplicacion__ciclo=self.ciclo)


@override_settings(INFORME_TOC_ESTRICTO=True)
class TestTocActualizado(FixturePruebaMixin, TestCase):
    """El DOCX descargado llega con el índice REAL (LibreOffice)."""

    @unittest.skipUnless(HAY_LIBREOFFICE, 'requiere LibreOffice (soffice)')
    def test_indice_poblado_con_paginas_reales(self):
        res = self._descargar_informe()
        self.assertEqual(res.status_code, 200)
        with zipfile.ZipFile(io.BytesIO(res.content)) as z:
            xml = z.read('word/document.xml').decode('utf-8')
        self.assertNotIn('Marcador no definido', xml)
        self.assertNotIn('Bookmark not defined', xml)

        doc = Document(io.BytesIO(res.content))
        entradas = [p.text.strip() for p in doc.paragraphs
                    if p.style is not None and p.style.name.lower().startswith('toc')
                    and p.text.strip()]
        self.assertGreaterEqual(len(entradas), len(ESTRUCTURA_CANONICA) - 2)

        textos_toc = '\n'.join(entradas)
        for canonico in TEXTOS_CANONICOS:
            self.assertIn(canonico.split('\t')[0][:40], textos_toc,
                          f'Entrada canónica ausente del índice: {canonico}')
        # La entrada retirada (9.4. Dimensiones) no debe regresar.
        self.assertNotIn('9.4. Dimensiones', textos_toc)

        paginas = []
        for e in entradas:
            m = re.search(r'(\d+)\s*$', e)
            if m:
                paginas.append(int(m.group(1)))
        self.assertGreaterEqual(len(paginas), 30, f'Índice sin páginas: {entradas[:5]}')
        self.assertTrue(all(p > 0 for p in paginas))
        self.assertGreater(len(set(paginas)), 5,
                           'Todas las páginas del índice son iguales (índice roto)')
        self.assertEqual(paginas, sorted(paginas), 'Páginas del índice desordenadas')


@override_settings(INFORME_TOC_ESTRICTO=False)
class TestSinBypassPorNombre(FixturePruebaMixin, TestCase):
    """Ningún tenant recibe archivos fijos por su nombre."""

    def test_tenants_con_bypass_historico_reciben_informe_dinamico(self):
        for nombre in ('San Luis', 'Zapotitlan', 'Saltillo'):
            self.tenant.nombre = nombre
            self.tenant.save(update_fields=['nombre'])
            try:
                res = self._descargar_informe()
                self.assertEqual(res.status_code, 200)
                self.assertEqual(res['Content-Type'], MIME_DOCX)
                nombre_archivo = f'Informe_Diagnostico_NOM035_{nombre.replace(" ", "_")}_2026.docx'
                self.assertIn(nombre_archivo, res['Content-Disposition'])
                doc = Document(io.BytesIO(res.content))
                # El contenido es el del generador dinámico (estructura
                # canónica), no el archivo fijo histórico.
                self.assertEqual(_encabezados(doc), ESTRUCTURA_CANONICA)
            finally:
                self.tenant.nombre = 'PRUEBA'
                self.tenant.save(update_fields=['nombre'])

    def test_codigo_sin_atajos_por_nombre(self):
        fuente = Path(__file__).with_name('views.py').read_text(encoding='utf-8')
        self.assertNotIn("== 'san luis'", fuente.lower())
        self.assertNotIn("== 'zapotitlan'", fuente.lower())
        self.assertNotIn("== 'saltillo'", fuente.lower())
        self.assertNotIn('PLANTA Zapot.pdf', fuente)


@override_settings(INFORME_TOC_ESTRICTO=False)
class TestRenderizadoSinDefectos(FixturePruebaMixin, TestCase):
    """Render LibreOffice → PDF: sin palabras partidas carácter por carácter,
    sin marcadores rotos y sin páginas vacías accidentales."""

    @unittest.skipUnless(HAY_LIBREOFFICE, 'requiere LibreOffice (soffice)')
    def test_render_pdf_limpio(self):
        res = self._descargar_informe()
        self.assertEqual(res.status_code, 200)
        with tempfile.TemporaryDirectory(prefix='nom035_render_') as tmp:
            tmp = Path(tmp)
            docx = tmp / 'informe.docx'
            docx.write_bytes(res.content)
            subprocess.run(
                [localizar_soffice(), '--headless', '--norestore',
                 f'-env:UserInstallation={(tmp / "perfil").as_uri()}',
                 '--convert-to', 'pdf', '--outdir', str(tmp), str(docx)],
                capture_output=True, timeout=300, check=False,
            )
            pdf = tmp / 'informe.pdf'
            self.assertTrue(pdf.exists(), 'LibreOffice no produjo el PDF')
            txt = tmp / 'informe.txt'
            pdftotext = None
            import shutil as _sh
            pdftotext = _sh.which('pdftotext')
            if not pdftotext:
                self.skipTest('pdftotext (Poppler) no disponible')
            subprocess.run([pdftotext, '-layout', str(pdf), str(txt)],
                           capture_output=True, timeout=120, check=True)
            texto = txt.read_text(encoding='utf-8', errors='replace')

        self.assertNotIn('Marcador no definido', texto)

        paginas = texto.split('\f')
        self.assertGreater(len(paginas), 20, 'El informe quedó demasiado corto')
        vacias = [i + 1 for i, p in enumerate(paginas[:-1]) if not p.strip()]
        self.assertEqual(vacias, [], f'Páginas vacías accidentales: {vacias}')

        # Palabras partidas carácter por carácter: 4+ líneas consecutivas de
        # un solo carácter alfabético (síntoma de columnas sin ancho).
        for num, pagina in enumerate(paginas, 1):
            consecutivas = 0
            for linea in pagina.split('\n'):
                if re.fullmatch(r'\s*[A-Za-zÁÉÍÓÚÑáéíóúñ]\s*', linea or ''):
                    consecutivas += 1
                    if consecutivas >= 4:
                        self.fail(f'Texto partido carácter por carácter en la página {num}')
                else:
                    consecutivas = 0


class TestBloqueosEstructurales(SimpleTestCase):
    """Un informe con catálogos incompletos NO puede emitirse."""

    def _componer(self, **kw):
        return componer_report_data(_raw([_persona('A', 'medio')] * 6, **kw),
                                    umbral_conf=5, umbral_faltantes_pct=10.0)

    def test_categorias_incompletas_bloquean(self):
        cats = [_cat(c, 'medio') for c in _CATEGORIAS_OFICIALES[:-1]] * 6
        rd = self._componer(categorias_ind=cats)
        self.assertFalse(rd['validaciones']['puede_emitirse'])
        self.assertTrue(any('categorias_oficiales_completas' in e
                            for e in rd['validaciones']['errores_criticos']))

    def test_dominios_incompletos_bloquean(self):
        doms = [_dom(d, 'medio') for d in _DOMINIOS_OFICIALES[:-1]] * 6
        rd = self._componer(dominios_ind=doms)
        self.assertFalse(rd['validaciones']['puede_emitirse'])
        self.assertTrue(any('dominios_oficiales_completos' in e
                            for e in rd['validaciones']['errores_criticos']))

    def test_listas_vacias_no_pasan(self):
        rd = self._componer(categorias_ind=[], dominios_ind=[])
        self.assertFalse(rd['validaciones']['puede_emitirse'])

    def test_dimensiones_incompletas_bloquean(self):
        raw = _raw([_persona('A', 'medio')] * 6)
        raw['dimensiones_observadas'] = ['Carga mental']  # 1 de 25
        rd = componer_report_data(raw, umbral_conf=5, umbral_faltantes_pct=10.0)
        self.assertFalse(rd['validaciones']['puede_emitirse'])
        self.assertTrue(any('dimensiones_oficiales_completas' in e
                            for e in rd['validaciones']['errores_criticos']))

    def test_dimensiones_condicionales_pueden_faltar(self):
        from documents.report_data import _DIMENSIONES_SIEMPRE
        raw = _raw([_persona('A', 'medio')] * 6)
        raw['dimensiones_observadas'] = sorted(_DIMENSIONES_SIEMPRE)  # 23
        rd = componer_report_data(raw, umbral_conf=5, umbral_faltantes_pct=10.0)
        self.assertTrue(rd['validaciones']['puede_emitirse'],
                        rd['validaciones']['errores_criticos'])
